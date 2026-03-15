import json
from docx import Document

# Dosya adını ve çıkış json dosyasını belirtin
DOCX_PATH = 'backend/Eğitim Katılım Formu 1-2.docx'
JSON_PATH = 'backend/egitim_katilim_formu_1_2.json'

def extract_docx_to_json(docx_path):
    doc = Document(docx_path)
    data = {
        'paragraphs': [],
        'tables': []
    }
    # Paragrafları ekle
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            data['paragraphs'].append(text)
    # Tabloları ekle
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        data['tables'].append(table_data)
    return data

def main():
    data = extract_docx_to_json(DOCX_PATH)
    with open(JSON_PATH, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"Word içeriği başarıyla {JSON_PATH} dosyasına aktarıldı.")

if __name__ == '__main__':
    main()
