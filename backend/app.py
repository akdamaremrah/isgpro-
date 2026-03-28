import os
import time
import json
import re
import concurrent.futures
import google.generativeai as genai
from dotenv import load_dotenv

load_dotenv()  # .env dosyasını en başta yükle

from flask import Flask, jsonify, request, send_from_directory, send_file
from werkzeug.utils import secure_filename
from flask_cors import CORS
from models import (
    db, Company, CompanyProfessional, User, Personnel, GenderEnum,
    RiskAssessment, RiskStatusEnum, calculate_service_minutes, Assignment,
    AssignmentTypeEnum, RoleEnum, NaceCode, DangerClass, City, District,
    HealthStatusEnum, DocumentTypeEnum, TrainingDocument, TrainingDocTypeEnum,
    HealthDocument, HealthDocTypeEnum
)
from datetime import datetime
import pandas as pd
import io
import requests
from sqlalchemy import func
from services.webhook_emitter import emit_event

app = Flask(__name__)

# ── CORS ──────────────────────────────────────────────────────────────────────
# Üretimde FRONTEND_URL env değişkeninden izin verilen origin alınır.
# Lokalde veya değişken yoksa tüm originlere izin verilir.
_frontend_url = os.environ.get('FRONTEND_URL', '')
if _frontend_url:
    CORS(app, resources={r"/api/*": {"origins": [_frontend_url, "http://localhost:5173"]}})
else:
    CORS(app)  # local development: allow all

def clean_tc(val):
    """Normalize TC kimlik no: strip '.0' suffix, scientific notation, ensure plain digit string."""
    if val is None:
        return None
    s = str(val).strip()
    if s.endswith('.0'):
        s = s[:-2]
    if 'e+' in s.lower():
        try:
            s = str(int(float(s)))
        except (ValueError, OverflowError):
            pass
    return s if s and s != 'nan' else None

# Configuration
basedir = os.path.abspath(os.path.dirname(__file__))
UPLOAD_FOLDER = os.path.join(basedir, 'uploads', 'personnel')
TRAINING_DOCS_FOLDER = os.path.join(basedir, 'uploads', 'training_docs')
HEALTH_DOCS_FOLDER = os.path.join(basedir, 'uploads', 'health_docs')
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['TRAINING_DOCS_FOLDER'] = TRAINING_DOCS_FOLDER
app.config['HEALTH_DOCS_FOLDER'] = HEALTH_DOCS_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload

# ── Database ──────────────────────────────────────────────────────────────────
# Render PostgreSQL: DATABASE_URL env değişkeni otomatik set edilir.
# Render, postgres:// şeması kullanır; SQLAlchemy postgresql:// ister → düzelt.
_db_url = os.environ.get('DATABASE_URL', '')
if _db_url.startswith('postgres://'):
    _db_url = _db_url.replace('postgres://', 'postgresql://', 1)
app.config['SQLALCHEMY_DATABASE_URI'] = _db_url or ('sqlite:///' + os.path.join(basedir, 'isg.db'))
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# ── Public base URL (fotoğraf linkleri için) ──────────────────────────────────
BASE_URL = os.environ.get('BASE_URL', 'http://localhost:5000')

db.init_app(app)

def sync_company_professionals(company_id: int):
    """
    Recalculates monthly_service_minutes for all professionals of a company
    based on the current personnel count.
    """
    c = Company.query.get(company_id)
    if not c or not c.nace_code or not c.nace_code.danger_class:
        return
    
    # ensure we have the most up-to-date count
    emp_count = Personnel.query.filter_by(company_id=company_id).count()
    hazard = c.nace_code.danger_class.name
    
    profs = CompanyProfessional.query.filter_by(company_id=company_id).all()
    for prof in profs:
        prof.monthly_service_minutes = calculate_service_minutes(hazard, emp_count, prof.role.name)
    
    db.session.commit()

@app.route('/api/uploads/personnel/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

# --- USER ROUTES ---
@app.route('/api/users', methods=['GET'])
def get_users():
    users = User.query.all()
    return jsonify([u.to_dict() for u in users]), 200

@app.route('/api/users', methods=['POST'])
def create_user():
    data = request.json
    try:
        new_user = User(
            full_name=data['full_name'],
            email=data['email'],
            password=data['password'],
            role=data.get('role', 'İSG Uzmanı')
        )
        db.session.add(new_user)
        db.session.commit()
        return jsonify(new_user.to_dict()), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

@app.route('/api/users/<int:id>', methods=['DELETE'])
def delete_user(id):
    u = User.query.get_or_404(id)
    try:
        db.session.delete(u)
        db.session.commit()
        return jsonify({'message': 'Kullanıcı silindi'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

# --- ROUTES ---
@app.route('/api/companies', methods=['GET'])
def get_companies():
    companies = Company.query.order_by(Company.official_title.asc()).all()
    result = []
    for c in companies:
        title = ' '.join((c.official_title or '').split())
        result.append({
            'id': c.id,
            'unvan': title,
            'tehlikeSinifi': c.nace_code.danger_class.name if c.nace_code and c.nace_code.danger_class else 'Tehlikeli',
            'sgkSicilNo': c.sgk_no,
            'calisanSayisi': c.employee_count,
            'isveren': c.employer_name,
            'isgUzmani': c.authorized_person or '',
            'durum': 'Aktif' if c.is_active else 'Pasif',
            # Additional details
            'short_name': c.short_name,
            'nace_code_id': c.nace_code_id,
            'city_id': c.city_id,
            'district_id': c.district_id,
            'address': c.address,
            'employer_role': c.employer_role,
            'employer_phone': c.employer_phone,
            'employer_email': c.employer_email,
            'authorized_person': c.authorized_person,
            'authorized_role': c.authorized_role,
            'authorized_phone': c.authorized_phone,
            'authorized_email': c.authorized_email
        })
    return jsonify(result), 200

@app.route('/api/companies/<int:id>', methods=['GET'])
def get_company(id):
    c = Company.query.get_or_404(id)
    return jsonify({
        'id': c.id,
        'short_name': c.short_name,
        'official_title': c.official_title,
        'sgk_no': c.sgk_no,
        'tehlikeSinifi': c.nace_code.danger_class.name if c.nace_code and c.nace_code.danger_class else 'Tehlikeli',
        'nace_code_id': c.nace_code_id,
        'nace_code': c.nace_code.code if c.nace_code else '',
        'nace_description': c.nace_code.description if c.nace_code else '',
        'employee_count': c.employee_count,
        'employer_name': c.employer_name,
        'employer_role': c.employer_role,
        'employer_phone': c.employer_phone,
        'employer_email': c.employer_email,
        'authorized_person': c.authorized_person,
        'authorized_role': c.authorized_role,
        'authorized_phone': c.authorized_phone,
        'authorized_email': c.authorized_email,
        'city_id': c.city_id,
        'district_id': c.district_id,
        'address': c.address,
        'is_active': c.is_active,
        'has_isg_kurulu': c.has_isg_kurulu
    }), 200

@app.route('/api/companies', methods=['POST'])
def create_company():
    data = request.json
    try:
        n_id = data.get('nace_code_id', 1)

        new_company = Company(
            short_name=data.get('short_name', data.get('official_title')),
            official_title=data.get('official_title'),
            sgk_no=data.get('sgk_no'),
            nace_code_id=n_id,
            employee_count=data.get('employee_count', 0),
            city_id=data.get('city_id', 34),
            district_id=data.get('district_id', 1),
            address=data.get('address', ''),
            
            employer_name=data.get('employer_name', ''),
            employer_role=data.get('employer_role', ''),
            employer_phone=data.get('employer_phone', ''),
            employer_email=data.get('employer_email', ''),
            
            authorized_person=data.get('authorized_person', ''),
            authorized_role=data.get('authorized_role', ''),
            authorized_phone=data.get('authorized_phone', ''),
            authorized_email=data.get('authorized_email', ''),
            
            is_active=True
        )
        db.session.add(new_company)
        db.session.commit()
        
        # EMIT EVENT
        emit_event("COMPANY_CREATED", {
            "id": new_company.id,
            "unvan": new_company.official_title,
            "sgk_no": new_company.sgk_no
        })
        
        return jsonify({'message': 'Firma başarıyla oluşturuldu', 'id': new_company.id}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

import io
from flask import send_file

@app.route('/api/companies/download-template', methods=['GET'])
def download_template():
    import pandas as pd
    
    # Define EXACT headers requested by user
    columns = [
        "Firma Adı/ Unvanı",
        "SGK Sicil No",
        "Tehlike Sınıfı",
        "Çalışan Sayısı",
        "İş Güvenliği Uzmanı Adı Soyadı",
        "İşyeri Hekimi Adı Soyadı",
        "Sözleşme Tarihleri"
    ]
    df = pd.DataFrame(columns=columns)
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Firmalar')
    
    output.seek(0)
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name='isg_firma_sablonu.xlsx'
    )

@app.route('/api/companies/bulk-upload', methods=['POST'])
def bulk_upload():
    if 'file' not in request.files:
        return jsonify({'error': 'Dosya bulunamadı'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Dosya seçilmedi'}), 400

    import pandas as pd
    try:
        # Read without headers to find the actual header row
        if file.filename.endswith('.csv'):
            df_raw = pd.read_csv(file, header=None, on_bad_lines='skip')
        else:
            df_raw = pd.read_excel(file, header=None)
            
        header_idx = 0
        for i, row in df_raw.iterrows():
            row_str = " ".join([str(x).lower() for x in row.values])
            if 'unvan' in row_str or 'sicil' in row_str or 'sgk' in row_str:
                header_idx = i
                break
                
        file.seek(0)
        if file.filename.endswith('.csv'):
            df = pd.read_csv(file, header=header_idx, on_bad_lines='skip')
        else:
            df = pd.read_excel(file, header=header_idx)

        unvan_col = next((c for c in df.columns if 'unvan' in str(c).lower()), None)
        sgk_col = next((c for c in df.columns if 'sicil' in str(c).lower() or 'sgk' in str(c).lower()), None)
        emp_col = next((c for c in df.columns if 'çalışan' in str(c).lower() or 'calisan' in str(c).lower()), None)
        tehlike_col = next((c for c in df.columns if 'tehlike' in str(c).lower()), None)
        igu_col = next((c for c in df.columns if 'uzman' in str(c).lower() or 'igu' in str(c).lower() or 'güvenliği' in str(c).lower()), None)
        iyh_col = next((c for c in df.columns if 'hekim' in str(c).lower() or 'iyh' in str(c).lower()), None)
        date_col = next((c for c in df.columns if 'tarih' in str(c).lower()), None)

        if not unvan_col or not sgk_col:
            return jsonify({'error': 'Unvan veya SGK Sicil No sütunları bulunamadı. Lütfen dosya formatını veya şablonu kontrol edin.'}), 400
            
        success = 0
        errors = 0
        
        for _, row in df.iterrows():
            unvan = str(row[unvan_col]).strip() if pd.notna(row[unvan_col]) else ''
            sgk = str(row[sgk_col]).strip() if pd.notna(row[sgk_col]) else ''
            
            if not unvan or not sgk or sgk == 'nan' or unvan == 'nan':
                errors += 1
                continue
                
            existing = Company.query.filter_by(sgk_no=sgk).first()
            if existing:
                errors += 1
                continue
                
            emp_count = 0
            if emp_col and pd.notna(row[emp_col]):
                try:
                    emp_count = int(float(row[emp_col]))
                except:
                    pass
                    
            tehlike_str = str(row[tehlike_col]).strip() if tehlike_col and pd.notna(row[tehlike_col]) else ''
            # Find a NACE code matching the danger class
            nace_id = 1
            if tehlike_str:
                dc = DangerClass.query.filter_by(name=tehlike_str).first()
                if dc:
                    nc = NaceCode.query.filter_by(danger_class_id=dc.id).first()
                    if nc:
                        nace_id = nc.id
                    
            db.session.begin_nested() # Create a savepoint for this row
            try:
                c = Company(
                    short_name=unvan[:100],
                    official_title=unvan,
                    sgk_no=sgk,
                    nace_code_id=nace_id,
                    employee_count=emp_count,
                    city_id=34,
                    district_id=1,
                    address='',
                    employer_name='',
                    authorized_person='',
                    is_active=True
                )
                db.session.add(c)
                db.session.flush() # Flush to get company ID

                date_val = str(row[date_col]).strip() if date_col and pd.notna(row[date_col]) else None
                
                # Check for IGU
                igu_val = str(row[igu_col]).strip() if igu_col and pd.notna(row[igu_col]) else ''
                if igu_val and igu_val != 'nan':
                    igu_prof = CompanyProfessional(
                        company_id=c.id,
                        professional_name=igu_val,
                        role='IGU',
                        certificate_class='A', # Default
                        contract_start_date=date_val,
                        is_active=True
                    )
                    db.session.add(igu_prof)
                
                # Check for IYH
                iyh_val = str(row[iyh_col]).strip() if iyh_col and pd.notna(row[iyh_col]) else ''
                if iyh_val and iyh_val != 'nan':
                    iyh_prof = CompanyProfessional(
                        company_id=c.id,
                        professional_name=iyh_val,
                        role='IYH',
                        certificate_class='Diğer',
                        contract_start_date=date_val,
                        is_active=True
                    )
                    db.session.add(iyh_prof)
                
                db.session.commit() # Commit the nested transaction
                success += 1
            except Exception as row_e:
                db.session.rollback() # Rollback to savepoint on row failure
                errors += 1
                print(f"Row Error for {unvan}: {str(row_e)}")
            
        # EMIT EVENT
        emit_event("BULK_UPLOAD_COMPLETED", {
            "success_count": success,
            "error_count": errors
        })
            
        return jsonify({'success_count': success, 'error_count': errors}), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Dosya okuma hatası: {str(e)}'}), 400

@app.route('/api/companies/<int:id>', methods=['PUT'])
def update_company(id):
    c = Company.query.get_or_404(id)
    data = request.json
    try:
        if 'short_name' in data: c.short_name = data['short_name']
        if 'official_title' in data: c.official_title = data['official_title']
        if 'sgk_no' in data: c.sgk_no = data['sgk_no']
        if 'nace_code_id' in data: c.nace_code_id = int(data['nace_code_id'])

        if 'employee_count' in data: c.employee_count = int(data['employee_count'])
        if 'city_id' in data: c.city_id = int(data['city_id'])
        if 'district_id' in data: c.district_id = int(data['district_id'])
        if 'address' in data: c.address = data['address']
        
        if 'employer_name' in data: c.employer_name = data['employer_name']
        if 'employer_role' in data: c.employer_role = data['employer_role']
        if 'employer_phone' in data: c.employer_phone = data['employer_phone']
        if 'employer_email' in data: c.employer_email = data['employer_email']
        
        if 'authorized_person' in data: c.authorized_person = data['authorized_person']
        if 'authorized_role' in data: c.authorized_role = data['authorized_role']
        if 'authorized_phone' in data: c.authorized_phone = data['authorized_phone']
        if 'authorized_email' in data: c.authorized_email = data['authorized_email']
        if 'has_isg_kurulu' in data: c.has_isg_kurulu = bool(data['has_isg_kurulu'])
        
        db.session.commit()
        # TRIGGER SYNC for Risk Team
        sync_risk_assessment_team(id)
        
        # EMIT EVENT
        emit_event("COMPANY_UPDATED", {
            "id": c.id,
            "unvan": c.official_title
        })
        
        return jsonify({'message': 'Firma başarıyla güncellendi', 'id': c.id}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

@app.route('/api/companies/<int:id>/suspend', methods=['PUT'])
def suspend_company(id):
    c = Company.query.get_or_404(id)
    try:
        c.is_active = not c.is_active
        db.session.commit()
        status = "Aktif" if c.is_active else "Pasif"
        
        # EMIT EVENT
        emit_event("COMPANY_STATUS_CHANGED", {
            "id": c.id,
            "status": status
        })
        
        return jsonify({'message': f'Firma durumu {status} olarak değiştirildi.'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

@app.route('/api/companies/<int:id>', methods=['DELETE'])
def delete_company(id):
    c = Company.query.get_or_404(id)
    try:
        deleted_id = c.id
        db.session.delete(c)
        db.session.commit()
        
        # EMIT EVENT
        emit_event("COMPANY_DELETED", {
            "id": deleted_id
        })
        
        return jsonify({'message': 'Firma silindi.'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

@app.route('/api/companies/bulk', methods=['DELETE'])
def bulk_delete_companies():
    data = request.json
    ids_to_delete = data.get('ids', [])
    if not ids_to_delete:
        return jsonify({'error': 'Silinecek firma IDleri belirtilmedi'}), 400
        
    try:
        companies = Company.query.filter(Company.id.in_(ids_to_delete)).all()
        for c in companies:
            db.session.delete(c)
        db.session.commit()
        
        # EMIT EVENT
        emit_event("BULK_DELETE_COMPLETED", {
            "count": len(companies),
            "ids": ids_to_delete
        })
        
        return jsonify({'message': f'{len(companies)} firma başarıyla silindi.'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

@app.route('/api/n8n/callback', methods=['POST'])
def n8n_callback():
    """
    Endpoint for n8n to send results back (e.g., document URLs, compliance status).
    """
    data = request.json
    event = data.get('event')
    payload = data.get('data', {})
    
    if event == 'DOCUMENT_GENERATED':
        # Example: Update a risk assessment with its PDF URL
        ra_id = payload.get('risk_assessment_id')
        pdf_url = payload.get('pdf_url')
        # Logic to update DB...
        
    elif event == 'COMPLIANCE_CHECK_RESULT':
        # Example: Update company compliance status
        pass

    return jsonify({'status': 'ok'}), 200

def seed_data():
    from models import DangerClass, NaceCode, City, District
    
    # 1. Danger Classes
    if not DangerClass.query.first():
        classes = [
            DangerClass(id=1, name='Tehlikeli', color_code='#f59e0b'),
            DangerClass(id=2, name='Çok Tehlikeli', color_code='#ef4444'),
            DangerClass(id=3, name='Az Tehlikeli', color_code='#10b981')
        ]
        for c in classes:
            db.session.merge(c)
        db.session.commit()
        print("Danger classes seeded.")

    # 2. Nace Codes
    if not NaceCode.query.first():
        codes = [
            NaceCode(id=1, code='GENEL-T', description='Tehlikeli İşler Genel', danger_class_id=1),
            NaceCode(id=2, code='GENEL-CT', description='Çok Tehlikeli İşler Genel', danger_class_id=2),
            NaceCode(id=3, code='GENEL-AT', description='Az Tehlikeli İşler Genel', danger_class_id=3)
        ]
        for nc in codes:
            db.session.merge(nc)
        db.session.commit()
        print("Nace codes seeded.")

    # 3. Cities & Districts (Essential for Company creation)
    if not City.query.get(34):
        istanbul = City(id=34, name='İstanbul')
        db.session.merge(istanbul)
        db.session.commit()
        if not District.query.get(1):
            d1 = District(id=1, city_id=34, name='Merkez')
            db.session.merge(d1)
            db.session.commit()
        print("City/District seeded.")

def setup_db():
    with app.app_context():
        db.create_all()
        seed_data()

# --- PROFESSIONALS ROUTES ---
@app.route('/api/professionals/directory', methods=['GET'])
def professionals_directory():
    """Return unique professionals across all companies, deduplicated by tc_kimlik_no."""
    profs = CompanyProfessional.query.all()
    seen = {}
    for p in profs:
        tc = clean_tc(p.tc_kimlik_no)
        key = tc if tc else f'_id_{p.id}'
        if key not in seen or (p.is_active and not seen[key].get('is_active')):
            seen[key] = {
                'id': p.id,
                'professional_name': p.professional_name,
                'role': p.role.value,
                'tc_kimlik_no': tc,
                'certificate_class': p.certificate_class,
                'certificate_number': p.certificate_number,
                'phone_number': p.phone_number,
                'email': p.email,
                'photoUrl': f"{BASE_URL}/api/uploads/personnel/{p.profile_photo}" if p.profile_photo else None
            }
    return jsonify(list(seen.values())), 200

@app.route('/api/companies/<int:company_id>/professionals', methods=['GET'])
def get_professionals(company_id):
    Company.query.get_or_404(company_id)
    profs = CompanyProfessional.query.filter_by(company_id=company_id).all()
    result = []
    for p in profs:
        result.append({
            'id': p.id,
            'role': p.role.value,
            'professional_name': p.professional_name,
            'contract_start_date': p.contract_start_date.isoformat() if p.contract_start_date else None,
            'is_active': p.is_active,
            'monthly_service_minutes': p.monthly_service_minutes,
            'tc_kimlik_no': clean_tc(p.tc_kimlik_no),
            'certificate_class': p.certificate_class,
            'certificate_number': p.certificate_number,
            'phone_number': p.phone_number,
            'email': p.email,
            'photoUrl': f"{BASE_URL}/api/uploads/personnel/{p.profile_photo}" if p.profile_photo else None
        })
    return jsonify(result), 200

@app.route('/api/companies/<int:company_id>/professionals', methods=['POST'])
def create_professional(company_id):
    c = Company.query.get_or_404(company_id)
    data = request.json
    try:
        from models import RoleEnum
        import datetime
        role_str = data.get('role', 'IGU')
        role = RoleEnum[role_str]
        start_date = datetime.date.fromisoformat(data.get('contract_start_date'))
        
        # Auto-calculate service minutes from DB hazard
        hazard = c.nace_code.danger_class.name if c.nace_code and c.nace_code.danger_class else 'Tehlikeli'
        emp_count = c.employee_count or 0
        minutes = calculate_service_minutes(hazard, emp_count, role_str)
        
        prof = CompanyProfessional(
            company_id=company_id,
            professional_name=data.get('professional_name', ''),
            role=role,
            contract_start_date=start_date,
            is_active=True,
            monthly_service_minutes=minutes,
            tc_kimlik_no=clean_tc(data.get('tc_kimlik_no')),
            certificate_class=data.get('certificate_class'),
            certificate_number=data.get('certificate_number'),
            phone_number=data.get('phone_number'),
            email=data.get('email')
        )
        db.session.add(prof)
        db.session.commit()
        
        # TRIGGER SYNC for Risk Team
        sync_risk_assessment_team(company_id)
        
        return jsonify({'message': 'Profesyonel eklendi', 'id': prof.id, 'monthly_service_minutes': minutes}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

@app.route('/api/professionals/<int:prof_id>', methods=['PUT'])
def update_professional(prof_id):
    prof = CompanyProfessional.query.get_or_404(prof_id)
    data = request.json
    try:
        from models import RoleEnum
        import datetime
        
        if 'professional_name' in data:
            prof.professional_name = data['professional_name']
        if 'role' in data:
            role_str = data['role']
            prof.role = RoleEnum[role_str]
        if 'contract_start_date' in data and data['contract_start_date']:
            prof.contract_start_date = datetime.date.fromisoformat(data['contract_start_date'])
        if 'tc_kimlik_no' in data: prof.tc_kimlik_no = clean_tc(data['tc_kimlik_no'])
        if 'certificate_class' in data: prof.certificate_class = data['certificate_class']
        if 'certificate_number' in data: prof.certificate_number = data['certificate_number']
        if 'phone_number' in data: prof.phone_number = data['phone_number']
        if 'email' in data: prof.email = data['email']
            
        # Recalculate just in case role changed
        c = prof.company
        hazard = c.nace_code.danger_class.name if c.nace_code and c.nace_code.danger_class else 'Tehlikeli'
        emp_count = c.employee_count or 0
        prof.monthly_service_minutes = calculate_service_minutes(hazard, emp_count, prof.role.name)
        
        db.session.flush()

        # Sync personal info to other companies with same tc_kimlik_no
        tc = clean_tc(prof.tc_kimlik_no)
        if tc:
            others = CompanyProfessional.query.filter(
                CompanyProfessional.tc_kimlik_no == tc,
                CompanyProfessional.id != prof.id
            ).all()
            for o in others:
                o.professional_name = prof.professional_name
                o.certificate_class = prof.certificate_class
                o.certificate_number = prof.certificate_number
                o.phone_number = prof.phone_number
                o.email = prof.email
                # Recalculate service minutes for each company
                oc = o.company
                oh = oc.nace_code.danger_class.name if oc.nace_code and oc.nace_code.danger_class else 'Tehlikeli'
                o.monthly_service_minutes = calculate_service_minutes(oh, oc.employee_count or 0, o.role.name)

        db.session.commit()
        
        # TRIGGER SYNC for Risk Team
        sync_risk_assessment_team(prof.company_id)
        
        return jsonify({'message': 'Profesyonel güncellendi'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

@app.route('/api/professionals/<int:prof_id>/photo', methods=['POST'])
def update_professional_photo(prof_id):
    prof = CompanyProfessional.query.get_or_404(prof_id)
    if 'photo' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['photo']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
        
    if file:
        import uuid
        ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else 'jpg'
        filename = f"prof_{prof_id}_{uuid.uuid4().hex[:8]}.{ext}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        if not os.path.exists(app.config['UPLOAD_FOLDER']):
            os.makedirs(app.config['UPLOAD_FOLDER'])
            
        file.save(filepath)
        prof.profile_photo = filename
        db.session.commit()
        
        return jsonify({
            'message': 'Fotoğraf güncellendi', 
            'photoUrl': f"{BASE_URL}/api/uploads/personnel/{filename}"
        }), 200

@app.route('/api/professionals/<int:prof_id>', methods=['DELETE'])
def delete_professional(prof_id):
    prof = CompanyProfessional.query.get_or_404(prof_id)
    try:
        company_id = prof.company_id
        db.session.delete(prof)
        db.session.commit()
        
        # TRIGGER SYNC for Risk Team
        sync_risk_assessment_team(company_id)
        
        return jsonify({'message': 'Profesyonel silindi.'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

# --- PERSONNEL ROUTES ---
@app.route('/api/companies/<int:company_id>/personnel', methods=['GET'])
def get_company_personnel(company_id):
    Company.query.get_or_404(company_id)
    # Alphabetical sorting
    personnel = Personnel.query.filter_by(company_id=company_id).order_by(Personnel.first_name, Personnel.last_name).all()
    result = []
    for p in personnel:
        result.append({
            'id': p.id,
            'tc': clean_tc(p.tc_no),
            'adSoyad': f"{p.first_name} {p.last_name}",
            'firstName': p.first_name,
            'lastName': p.last_name,
            'unvan': p.job_title,
            'dogumTarihi': p.birth_date.isoformat() if p.birth_date else None,
            'cinsiyet': p.gender.value,
            'iseGiris': p.hire_date.isoformat() if p.hire_date else None,
            'gsm': p.phone_number,
            'bloodType': p.blood_type,
            'educationLevel': p.education_level,
            'maritalStatus': p.marital_status,
            'disabilityStatus': p.disability_status,
            'trainingDate1': p.training_date_1.isoformat() if p.training_date_1 else None,
            'trainingDate2': p.training_date_2.isoformat() if p.training_date_2 else None,
            'healthExamDate': p.health_exam_date.isoformat() if p.health_exam_date else None,
            'trainingValidityDate': p.training_validity_date.isoformat() if p.training_validity_date else None,
            'healthValidityDate': p.health_validity_date.isoformat() if p.health_validity_date else None,
            'egitimGecerlilik': p.training_validity_date.isoformat() if p.training_validity_date else '-',
            'saglikGecerlilik': p.health_validity_date.isoformat() if p.health_validity_date else '-',
            'competencies': p.professional_competencies or [],
            'photoUrl': f"{BASE_URL}/api/uploads/personnel/{p.photo_path}" if p.photo_path else None,
            'is_active': p.is_active
        })
    return jsonify(result), 200

@app.route('/api/companies/<int:company_id>/personnel', methods=['POST'])
def create_personnel(company_id):
    Company.query.get_or_404(company_id)
    data = request.json
    try:
        existing = Personnel.query.filter_by(company_id=company_id, tc_no=clean_tc(data['tc'])).first()
        if existing:
            return jsonify({'error': 'Bu TC kimlik numarasına sahip personel zaten kayıtlı.'}), 400

        new_p = Personnel(
            company_id=company_id,
            tc_no=clean_tc(data.get('tc')),
            first_name=data.get('firstName'),
            last_name=data.get('lastName'),
            job_title=data.get('unvan', 'Personel'),
            birth_date=datetime.strptime(data['dogumTarihi'], '%Y-%m-%d').date() if data.get('dogumTarihi') else None,
            gender=GenderEnum(data['cinsiyet']) if data.get('cinsiyet') else GenderEnum.E,
            hire_date=datetime.strptime(data['iseGiris'], '%Y-%m-%d').date() if data.get('iseGiris') else None,
            phone_number=data.get('gsm'),
            blood_type=data.get('bloodType'),
            education_level=data.get('educationLevel'),
            marital_status=data.get('maritalStatus'),
            disability_status=data.get('disabilityStatus', 'Yok'),
            training_date_1=datetime.strptime(data['trainingDate1'], '%Y-%m-%d').date() if data.get('trainingDate1') else None,
            training_date_2=datetime.strptime(data['trainingDate2'], '%Y-%m-%d').date() if data.get('trainingDate2') else None,
            health_exam_date=datetime.strptime(data['healthExamDate'], '%Y-%m-%d').date() if data.get('healthExamDate') else None,
            professional_competencies=data.get('competencies', []),
            is_active=True
        )
        new_p.calculate_validity_dates()
        db.session.add(new_p)
        db.session.commit()
        return jsonify({'message': 'Personel başarıyla eklendi', 'id': new_p.id}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

@app.route('/api/personnel/<int:id>', methods=['PUT'])
def update_personnel(id):
    p = Personnel.query.get_or_404(id)
    data = request.json
    try:
        if 'tc' in data: p.tc_no = clean_tc(data['tc'])
        if 'firstName' in data: p.first_name = data['firstName']
        if 'lastName' in data: p.last_name = data['lastName']
        if 'unvan' in data: p.job_title = data['unvan']
        if 'dogumTarihi' in data:
             p.birth_date = datetime.strptime(data['dogumTarihi'], '%Y-%m-%d').date() if data['dogumTarihi'] else None
        if 'cinsiyet' in data and data['cinsiyet']: p.gender = GenderEnum(data['cinsiyet'])
        if 'iseGiris' in data:
             p.hire_date = datetime.strptime(data['iseGiris'], '%Y-%m-%d').date() if data['iseGiris'] else None
        if 'gsm' in data: p.phone_number = data['gsm']
        
        # New fields
        if 'bloodType' in data: p.blood_type = data['bloodType']
        if 'educationLevel' in data: p.education_level = data['educationLevel']
        if 'maritalStatus' in data: p.marital_status = data['maritalStatus']
        if 'disabilityStatus' in data: p.disability_status = data['disabilityStatus']
        if 'trainingDate1' in data:
            p.training_date_1 = datetime.strptime(data['trainingDate1'], '%Y-%m-%d').date() if data['trainingDate1'] else None
        if 'trainingDate2' in data:
            p.training_date_2 = datetime.strptime(data['trainingDate2'], '%Y-%m-%d').date() if data['trainingDate2'] else None
        if 'healthExamDate' in data:
            p.health_exam_date = datetime.strptime(data['healthExamDate'], '%Y-%m-%d').date() if data['healthExamDate'] else None
        if 'competencies' in data:
            p.professional_competencies = data['competencies']
            
        p.calculate_validity_dates()
        db.session.commit()
        return jsonify({'message': 'Personel güncellendi'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

@app.route('/api/personnel/<int:id>', methods=['DELETE'])
def delete_personnel(id):
    p = Personnel.query.get_or_404(id)
    try:
        db.session.delete(p)
        db.session.commit()
        return jsonify({'message': 'Personel silindi.'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

@app.route('/api/personnel/bulk', methods=['DELETE'])
def bulk_delete_personnel():
    data = request.json
    ids_to_delete = data.get('ids', [])
    if not ids_to_delete:
        return jsonify({'error': 'Silinecek personel IDleri belirtilmedi'}), 400

    try:
        personnel_list = Personnel.query.filter(Personnel.id.in_(ids_to_delete)).all()
        company_ids = set(p.company_id for p in personnel_list)
        for p in personnel_list:
            db.session.delete(p)
        db.session.commit()

        for cid in company_ids:
            sync_company_professionals(cid)

        return jsonify({'message': f'{len(personnel_list)} personel başarıyla silindi.'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

# --- TRAINING DATE ROUTES ---
@app.route('/api/companies/<int:company_id>/personnel/training-dates', methods=['PUT'])
def bulk_update_training_dates(company_id):
    """Seçilen personellerin eğitim tarihlerini toplu günceller."""
    Company.query.get_or_404(company_id)
    data = request.json or {}
    personnel_ids = data.get('personnelIds', [])
    date1 = data.get('trainingDate1')
    date2 = data.get('trainingDate2')

    if not personnel_ids:
        return jsonify({'error': 'Personel seçilmedi'}), 400
    if not date1 and not date2:
        return jsonify({'error': 'En az bir eğitim tarihi girilmeli'}), 400

    personnel_list = Personnel.query.filter(
        Personnel.id.in_(personnel_ids),
        Personnel.company_id == company_id
    ).all()

    updated = 0
    for p in personnel_list:
        if date1:
            p.training_date_1 = datetime.strptime(date1, '%Y-%m-%d').date()
        if date2:
            p.training_date_2 = datetime.strptime(date2, '%Y-%m-%d').date()
        p.calculate_validity_dates()
        updated += 1

    db.session.commit()
    return jsonify({'message': f'{updated} personelin eğitim tarihleri güncellendi', 'updated': updated}), 200

@app.route('/api/companies/<int:company_id>/personnel/health-dates', methods=['PUT'])
def bulk_update_health_dates(company_id):
    """Seçilen personellerin sağlık muayene tarihini toplu günceller."""
    Company.query.get_or_404(company_id)
    data = request.json or {}
    personnel_ids = data.get('personnelIds', [])
    health_exam_date = data.get('healthExamDate')

    if not personnel_ids:
        return jsonify({'error': 'Personel seçilmedi'}), 400
    if not health_exam_date:
        return jsonify({'error': 'Muayene tarihi girilmeli'}), 400

    personnel_list = Personnel.query.filter(
        Personnel.id.in_(personnel_ids),
        Personnel.company_id == company_id
    ).all()

    updated = 0
    for p in personnel_list:
        p.health_exam_date = datetime.strptime(health_exam_date, '%Y-%m-%d').date()
        p.calculate_validity_dates()
        updated += 1

    db.session.commit()
    return jsonify({'message': f'{updated} personelin sağlık muayene tarihleri güncellendi', 'updated': updated}), 200

# --- ASSIGNMENT ROUTES ---
@app.route('/api/companies/<int:company_id>/assignments', methods=['GET'])
def get_assignments(company_id):
    assignments = Assignment.query.filter_by(company_id=company_id).all()
    result = []
    for a in assignments:
        p_active = True
        if a.professional_id and a.professional:
            p_active = a.professional.is_active
        elif a.personnel_id and a.personnel:
            p_active = a.personnel.is_active
            
        result.append({
            'id': a.id,
            'personnel_id': a.personnel_id,
            'professional_id': a.professional_id,
            'employer_id': a.employer_id,
            'assignment_type': a.assignment_type.name,
            'assignment_date': a.assignment_date.isoformat() if a.assignment_date else None,
            'training_date': a.training_date.isoformat() if a.training_date else None,
            'training_validity_date': a.training_validity_date.isoformat() if a.training_validity_date else None,
            'role_in_team': a.person_role,
            'personnel_name': a.person_name,
            'is_personnel_active': p_active
        })
    return jsonify(result), 200

@app.route('/api/assignments/<int:assignment_id>', methods=['PUT'])
def update_assignment(assignment_id):
    a = Assignment.query.get_or_404(assignment_id)
    data = request.json
    import datetime as dt
    try:
        if 'assignment_date' in data:
            a.assignment_date = dt.date.fromisoformat(data['assignment_date']) if data['assignment_date'] else a.assignment_date
        if 'training_date' in data:
            a.training_date = dt.date.fromisoformat(data['training_date']) if data['training_date'] else None
        if 'training_validity_date' in data:
            a.training_validity_date = dt.date.fromisoformat(data['training_validity_date']) if data['training_validity_date'] else None
        db.session.commit()
        return jsonify({'message': 'Güncellendi'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

def sync_risk_assessment_team(company_id):
    """
    KATI KURAL: Bu fonksiyon KESİNLİKLE yeni Personnel kaydı oluşturmaz.
    İşveren (employer_id), İGU/IYH (professional_id) ve Temsilciler (personnel_id)
    bilgilerini doğrudan ilgili tablolardan çeker.
    """
    try:
        company = Company.query.get(company_id)
        if not company:
            return

        # 1. Mevcut risk ekibi atamalarını çek
        existing_assignments = Assignment.query.filter_by(
            company_id=company_id, 
            assignment_type=AssignmentTypeEnum.Risk_Degerlendirme_Ekibi
        ).all()
        
        to_keep_ids = []
        
        def ensure_assignment(p_id=None, prof_id=None, emp_id=None, role=None):
            match = None
            for a in existing_assignments:
                if p_id and a.personnel_id == p_id: match = a
                elif prof_id and a.professional_id == prof_id: match = a
                elif emp_id and a.employer_id == emp_id: match = a
                if match: break
                
            if not match:
                match = Assignment(
                    company_id=company_id,
                    personnel_id=p_id,
                    professional_id=prof_id,
                    employer_id=emp_id,
                    assignment_type=AssignmentTypeEnum.Risk_Degerlendirme_Ekibi,
                    role_in_team=role
                )
                db.session.add(match)
                db.session.flush()
            else:
                if role and match.role_in_team != role:
                    match.role_in_team = role
            
            to_keep_ids.append(match.id)

        # 2. İşveren (Her zaman)
        ensure_assignment(emp_id=company.id, role="İşveren / İşveren Vekili")

        # 3. İGU / İYH (Aktif profesyoneller)
        active_profs = CompanyProfessional.query.filter_by(company_id=company_id, is_active=True).all()
        for prof in active_profs:
            if prof.role in [RoleEnum.IGU, RoleEnum.IYH]:
                ensure_assignment(prof_id=prof.id, role=prof.role.value)

        # 4. Çalışan Temsilcileri
        rep_assignments = Assignment.query.filter_by(
            company_id=company_id,
            assignment_type=AssignmentTypeEnum.Calisan_Temsilcisi
        ).all()
        for ra in rep_assignments:
            if ra.personnel_id and ra.personnel.is_active:
                ensure_assignment(p_id=ra.personnel_id, role="Çalışan Temsilcisi")

        # 5. Destek Elemanları
        support_assignments = Assignment.query.filter_by(
            company_id=company_id,
            assignment_type=AssignmentTypeEnum.Destek_Elemani
        ).all()
        for sa in support_assignments:
            if sa.personnel_id and sa.personnel.is_active:
                ensure_assignment(p_id=sa.personnel_id, role="Destek Elemanı")

        # 6. Fazlalıkları sil
        for a in existing_assignments:
            if a.id not in to_keep_ids:
                db.session.delete(a)

        db.session.commit()
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Sync error for company {company_id}: {e}")

@app.route('/api/companies/<int:company_id>/assignments/bulk', methods=['POST'])
def bulk_save_assignments(company_id):
    data = request.json
    new_assignments_list = data.get('assignments', [])
    
    try:
        # FULL SYNC: Delete ALL existing assignments for this company
        Assignment.query.filter_by(company_id=company_id).delete()
        
        # Add new ones from the authoritative payload
        for item in new_assignments_list:
            if not item.get('personnel_id'): continue
            new_a = Assignment(
                company_id=company_id,
                personnel_id=item['personnel_id'],
                assignment_type=AssignmentTypeEnum[item['type']],
                role_in_team=item.get('role')
            )
            db.session.add(new_a)
            
        db.session.commit()
        
        # TRIGGER SYNC
        sync_risk_assessment_team(company_id)
        
        return jsonify({'message': 'Atamalar başarıyla güncellendi'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

@app.route('/api/personnel/<int:id>', methods=['GET'])
def get_personnel_detail(id):
    p = Personnel.query.get_or_404(id)
    # Get assignments for this personnel
    assignments = Assignment.query.filter_by(personnel_id=p.id).all()
    roles = [a.assignment_type.value for a in assignments]
    # Deduplicate
    roles = list(dict.fromkeys(roles))
    return jsonify({
        'id': p.id,
        'tc': clean_tc(p.tc_no),
        'adSoyad': f"{p.first_name} {p.last_name}",
        'firstName': p.first_name,
        'lastName': p.last_name,
        'unvan': p.job_title,
        'dogumTarihi': p.birth_date.isoformat() if p.birth_date else None,
        'cinsiyet': p.gender.value,
        'iseGiris': p.hire_date.isoformat() if p.hire_date else None,
        'gsm': p.phone_number,
        'bloodType': p.blood_type,
        'educationLevel': p.education_level,
        'maritalStatus': p.marital_status,
        'disabilityStatus': p.disability_status,
        'trainingDate1': p.training_date_1.isoformat() if p.training_date_1 else None,
        'trainingDate2': p.training_date_2.isoformat() if p.training_date_2 else None,
        'healthExamDate': p.health_exam_date.isoformat() if p.health_exam_date else None,
        'trainingValidityDate': p.training_validity_date.isoformat() if p.training_validity_date else None,
        'healthValidityDate': p.health_validity_date.isoformat() if p.health_validity_date else None,
        'competencies': p.professional_competencies or [],
        'photoUrl': f"{BASE_URL}/api/uploads/personnel/{p.photo_path}" if p.photo_path else None,
        'is_active': p.is_active,
        'companyId': p.company_id,
        'assignments': roles
    }), 200

@app.route('/api/personnel/<int:id>/photo', methods=['POST'])
def upload_personnel_photo(id):
    p = Personnel.query.get_or_404(id)
    if 'file' not in request.files:
        return jsonify({'error': 'Dosya bulunamadı'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Dosya seçilmedi'}), 400
    
    if file:
        filename = secure_filename(f"personnel_{id}_{file.filename}")
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        p.photo_path = filename
        db.session.commit()
        return jsonify({
            'message': 'Fotoğraf yüklendi',
            'photoUrl': f"{BASE_URL}/api/uploads/personnel/{filename}"
        }), 200

@app.route('/api/personnel/download-template', methods=['GET'])
def download_personnel_template():
    columns = [
        "Ad Soyad", "TC Kimlik No", "Unvan", "Doğum Tarihi (YYYY-AA-GG)", 
        "Cinsiyet (E/K)", "İşe Giriş Tarihi (YYYY-AA-GG)", "GSM No",
        "Öğrenim Durumu", "Medeni Hali", "Kan Grubu", "Engel Durumu",
        "Eğitim Tarihi 1", "Eğitim Tarihi 2", "Sağlık Muayenesi Tarihi",
        "Mesleki Yeterlilik (Virgülle ayırın)"
    ]
    df = pd.DataFrame(columns=columns)
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Personel Listesi')
    output.seek(0)
    
    return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 
                     as_attachment=True, download_name="personel_yukleme_sablonu.xlsx")

@app.route('/api/companies/<int:company_id>/personnel/bulk-upload', methods=['POST'])
def bulk_upload_personnel(company_id):
    Company.query.get_or_404(company_id)
    if 'file' not in request.files:
        return jsonify({'error': 'Dosya bulunamadı'}), 400
    
    file = request.files['file']
    try:
        if file.filename.endswith('.csv'):
            df = pd.read_csv(file)
        else:
            df = pd.read_excel(file)
        
        # Mapping helpers
        name_col = next((c for c in df.columns if 'ad' in str(c).lower() and 'soyad' in str(c).lower()), 
                        next((c for c in df.columns if 'ad' in str(c).lower() or 'soyad' in str(c).lower()), None))
        tc_col = next((c for c in df.columns if 'tc' in str(c).lower() or 'kimlik' in str(c).lower()), None)
        job_col = next((c for c in df.columns if 'unvan' in str(c).lower() or 'görev' in str(c).lower()), None)
        birth_col = next((c for c in df.columns if 'doğum' in str(c).lower() or 'dogum' in str(c).lower()), None)
        gender_col = next((c for c in df.columns if 'cinsiyet' in str(c).lower()), None)
        hire_col = next((c for c in df.columns if 'iş' in str(c).lower() or 'giris' in str(c).lower()), None)
        gsm_col = next((c for c in df.columns if 'gsm' in str(c).lower() or 'tel' in str(c).lower()), None)
        
        edu_col = next((c for c in df.columns if 'öğrenim' in str(c).lower() or 'egitim' in str(c).lower()), None)
        marital_col = next((c for c in df.columns if 'medeni' in str(c).lower()), None)
        blood_col = next((c for c in df.columns if 'kan' in str(c).lower()), None)
        disability_col = next((c for c in df.columns if 'engel' in str(c).lower()), None)
        
        t1_col = next((c for c in df.columns if 'eğitim' in str(c).lower() and '1' in str(c).lower()), None)
        t2_col = next((c for c in df.columns if 'eğitim' in str(c).lower() and '2' in str(c).lower()), None)
        h_col = next((c for c in df.columns if 'sağlık' in str(c).lower() or 'muayene' in str(c).lower()), None)
        comp_col = next((c for c in df.columns if 'yeterlilik' in str(c).lower()), None)

        success = 0
        skipped = 0
        
        def clean_numeric_str(val):
            if pd.isna(val): return None
            s = str(val).strip()
            if s.endswith('.0'): s = s[:-2]
            if 'e+' in s.lower(): # scientific notation
                try: s = str(int(float(s)))
                except: pass
            return s if s and s != 'nan' else None

        for _, row in df.iterrows():
            tc = clean_numeric_str(row[tc_col]) if tc_col else None
            if not tc:
                skipped += 1
                continue
            
            # Check for existing in this company
            existing = Personnel.query.filter_by(company_id=company_id, tc_no=tc).first()
            
            name_parts = str(row[name_col]).strip().split(' ') if name_col and pd.notna(row[name_col]) else ["İsimsiz"]
            first_name = " ".join(name_parts[:-1]) if len(name_parts) > 1 else name_parts[0]
            last_name = name_parts[-1] if len(name_parts) > 1 else ""

            gender_val = str(row[gender_col]).strip().upper() if gender_col and pd.notna(row[gender_col]) else "E"
            if gender_val.startswith('K'): gender_val = "K"
            else: gender_val = "E"

            try:
                def parse_date(col):
                    if col and pd.notna(row[col]):
                        if isinstance(row[col], datetime): return row[col].date()
                        try: return pd.to_datetime(row[col]).date()
                        except: return None
                    return None

                birth_date = parse_date(birth_col)
                hire_date = parse_date(hire_col)
                t1_date = parse_date(t1_col)
                t2_date = parse_date(t2_col)
                h_date = parse_date(h_col)
                
                competencies = []
                if comp_col and pd.notna(row[comp_col]):
                    competencies = [s.strip() for s in str(row[comp_col]).split(',') if s.strip()]

                if existing:
                    existing.first_name = first_name
                    existing.last_name = last_name
                    if job_col and pd.notna(row[job_col]):
                        existing.job_title = str(row[job_col]).strip()
                    existing.birth_date = birth_date or existing.birth_date
                    existing.gender = GenderEnum(gender_val) if gender_val else existing.gender
                    existing.hire_date = hire_date or existing.hire_date
                    if gsm_col and pd.notna(row[gsm_col]):
                        existing.phone_number = clean_numeric_str(row[gsm_col])
                    if edu_col and pd.notna(row[edu_col]):
                        existing.education_level = str(row[edu_col]).strip()
                    if marital_col and pd.notna(row[marital_col]):
                        existing.marital_status = str(row[marital_col]).strip()
                    if blood_col and pd.notna(row[blood_col]):
                        existing.blood_type = str(row[blood_col]).strip()
                    if disability_col and pd.notna(row[disability_col]):
                        existing.disability_status = str(row[disability_col]).strip()
                    
                    existing.training_date_1 = t1_date or existing.training_date_1
                    existing.training_date_2 = t2_date or existing.training_date_2
                    existing.health_exam_date = h_date or existing.health_exam_date
                    if competencies:
                        existing.professional_competencies = competencies
                    existing.calculate_validity_dates()
                else:
                    new_p = Personnel(
                        company_id=company_id,
                        tc_no=tc,
                        first_name=first_name,
                        last_name=last_name,
                        job_title=str(row[job_col]).strip() if job_col and pd.notna(row[job_col]) else "Personel",
                        birth_date=birth_date,
                        gender=GenderEnum(gender_val) if gender_val else GenderEnum.E,
                        hire_date=hire_date,
                        phone_number=clean_numeric_str(row[gsm_col]) if gsm_col else None,
                        education_level=str(row[edu_col]).strip() if edu_col and pd.notna(row[edu_col]) else None,
                        marital_status=str(row[marital_col]).strip() if marital_col and pd.notna(row[marital_col]) else None,
                        blood_type=str(row[blood_col]).strip() if blood_col and pd.notna(row[blood_col]) else None,
                        disability_status=str(row[disability_col]).strip() if disability_col and pd.notna(row[disability_col]) else "Yok",
                        training_date_1=t1_date,
                        training_date_2=t2_date,
                        health_exam_date=h_date,
                        professional_competencies=competencies
                    )
                    new_p.calculate_validity_dates()
                    db.session.add(new_p)
                
                success += 1
            except Exception as e:
                print(f"Row skip error: {e}")
                skipped += 1
        
        db.session.commit()
        return jsonify({'success_count': success, 'skipped_count': skipped}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

# --- RISK ASSESSMENT ROUTES ---
@app.route('/api/companies/<int:company_id>/risks', methods=['GET'])
def get_company_risks(company_id):
    Company.query.get_or_404(company_id)
    min_score = request.args.get('min_score', type=float)
    filter_type = request.args.get('filter')
    
    query = RiskAssessment.query.filter_by(company_id=company_id)
    if min_score is not None:
        query = query.filter(RiskAssessment.initial_risk_score > min_score)
    
    risks = query.order_by(RiskAssessment.id).all()
    result = []
    for r in risks:
        result.append({
            'id': r.id,
            'activity': r.activity,
            'hazard': r.hazard,
            'riskEffect': r.risk_effect,
            'correctiveMeasures': r.corrective_measures,
            'existingMeasures': r.existing_measures,
            'initialProbability': r.initial_probability,
            'initialSeverity': r.initial_severity,
            'initialFrequency': r.initial_frequency,
            'initialRiskScore': r.initial_risk_score,
            'initialRiskLevel': r.initial_risk_level,
            'finalProbability': r.final_probability,
            'finalSeverity': r.final_severity,
            'finalFrequency': r.final_frequency,
            'finalRiskScore': r.final_risk_score,
            'finalRiskLevel': r.final_risk_level,
            'responsiblePerson': r.responsible_person,
            'deadline': r.deadline,
        })
    return jsonify(result), 200

@app.route('/api/companies/<int:company_id>/risks', methods=['POST'])
def create_risk(company_id):
    Company.query.get_or_404(company_id)
    data = request.json
    try:
        new_risk = RiskAssessment(
            company_id=company_id,
            activity=data.get('activity', ''),
            hazard=data.get('hazard', ''),
            risk_effect=data.get('riskEffect', ''),
            corrective_measures=data.get('correctiveMeasures'),
            existing_measures=data.get('existingMeasures'),
            initial_probability=float(data.get('initialProbability', 0)),
            initial_severity=float(data.get('initialSeverity', 0)),
            initial_frequency=float(data.get('initialFrequency', 0)),
            final_probability=float(data.get('finalProbability', 0)),
            final_severity=float(data.get('finalSeverity', 0)),
            final_frequency=float(data.get('finalFrequency', 0)),
            responsible_person=data.get('responsiblePerson'),
            deadline=data.get('deadline'),
        )
        new_risk.calculate_risks()
        db.session.add(new_risk)
        db.session.commit()
        return jsonify({'message': 'Risk kaydı oluşturuldu', 'id': new_risk.id}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

@app.route('/api/risks/<int:id>', methods=['PUT'])
def update_risk(id):
    r = RiskAssessment.query.get_or_404(id)
    data = request.json
    try:
        if 'activity' in data: r.activity = data['activity']
        if 'hazard' in data: r.hazard = data['hazard']
        if 'riskEffect' in data: r.risk_effect = data['riskEffect']
        if 'correctiveMeasures' in data: r.corrective_measures = data['correctiveMeasures']
        if 'existingMeasures' in data: r.existing_measures = data['existingMeasures']
        if 'initialProbability' in data: r.initial_probability = float(data['initialProbability'])
        if 'initialSeverity' in data: r.initial_severity = float(data['initialSeverity'])
        if 'initialFrequency' in data: r.initial_frequency = float(data['initialFrequency'])
        if 'finalProbability' in data: r.final_probability = float(data['finalProbability'])
        if 'finalSeverity' in data: r.final_severity = float(data['finalSeverity'])
        if 'finalFrequency' in data: r.final_frequency = float(data['finalFrequency'])
        if 'responsiblePerson' in data: r.responsible_person = data['responsiblePerson']
        if 'deadline' in data: r.deadline = data['deadline']
        r.calculate_risks()
        db.session.commit()
        return jsonify({'message': 'Risk güncellendi'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

@app.route('/api/risks/<int:id>', methods=['DELETE'])
def delete_risk(id):
    r = RiskAssessment.query.get_or_404(id)
    try:
        db.session.delete(r)
        db.session.commit()
        return jsonify({'message': 'Risk silindi.'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

@app.route('/api/companies/<int:company_id>/risks/export', methods=['GET'])
def export_risks_excel(company_id):
    company = Company.query.get_or_404(company_id)
    risks = RiskAssessment.query.filter_by(company_id=company_id).order_by(RiskAssessment.id).all()

    report_date = request.args.get('report_date', '')
    validity_date = request.args.get('validity_date', '')

    def fmt_date(d):
        if not d: return '—'
        try:
            parts = d.split('-')
            return f"{parts[2]}.{parts[1]}.{parts[0]}"
        except:
            return d

    company_title = company.official_title or '—'
    company_sgk = company.sgk_no or '—'
    company_address = company.address or '—'
    tehlike_sinifi = company.nace_code.danger_class.name if company.nace_code and company.nace_code.danger_class else '—'

    # Firma bilgileri satırları (Excel başına eklenecek)
    # Her satır: [Etiket, Değer] şeklinde 2 sütunlu blok, toplam 18 sütuna yayılacak
    firma_info = [
        ('Firma Unvanı', company_title),
        ('SGK Sicil No', company_sgk),
        ('Firma Adresi', company_address),
        ('Tehlike Sınıfı', tehlike_sinifi),
        ('Risk Değerlendirme Raporu Tarihi', fmt_date(report_date)),
        ('Geçerlilik Tarihi', fmt_date(validity_date)),
    ]
    # firma_info satır sayısı + 1 boş satır = DATA_START_ROW
    FIRMA_ROWS = len(firma_info) + 1  # +1 boşluk
    DATA_HEADER_ROW = FIRMA_ROWS      # tablo başlığı buradan başlar (0-indexed)
    DATA_START_ROW = FIRMA_ROWS + 2   # veri satırları buradan (header 2 satır)

    # Sütun Başlıkları (DataFrame için arka planda tutuyoruz, Excele biz yazacağız)
    columns = [
        'Sıra No', 
        'Yapılan İş / Faaliyet / Ekipman / Ortam', 
        'Tehlike', 
        'Risk / Olası Etki', 
        'Mevcutta Alınan Önlemler', 
        'İlk Risk - Olasılık', 'İlk Risk - Frekans', 'İlk Risk - Şiddet', 'İlk Risk - Risk Puanı', 'İlk Risk - Risk Düzeyi',
        'Açıklama / Düzeltici ve Önleyici Tedbirler', 
        'Son Risk - Olasılık', 'Son Risk - Frekans', 'Son Risk - Şiddet', 'Son Risk - Risk Puanı', 'Son Risk - Risk Düzeyi',
        'Sorumlu', 'Termin Süresi'
    ]
    
    rows = []
    for i, r in enumerate(risks, 1):
        rows.append([
            i,
            r.activity,
            r.hazard,
            r.risk_effect,
            r.existing_measures or '',
            r.initial_probability,
            r.initial_frequency,
            r.initial_severity,
            r.initial_risk_score,
            r.initial_risk_level,
            r.corrective_measures or '',
            r.final_probability,
            r.final_frequency,
            r.final_severity,
            r.final_risk_score,
            r.final_risk_level,
            r.responsible_person or '',
            r.deadline or ''
        ])
    
    df = pd.DataFrame(rows, columns=columns)
    output = io.BytesIO()

    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, header=False, startrow=DATA_START_ROW, sheet_name='Risk Değerlendirmesi')

        workbook  = writer.book
        worksheet = writer.sheets['Risk Değerlendirmesi']

        # -----------------------------
        # STİL (FORMAT) TANIMLAMALARI
        # -----------------------------
        base_format = workbook.add_format({
            'align': 'center', 'valign': 'vcenter', 'border': 1
        })
        wrap_format = workbook.add_format({
            'align': 'left', 'valign': 'vcenter', 'border': 1, 'text_wrap': True
        })
        header_format = workbook.add_format({
            'bold': True, 'align': 'center', 'valign': 'vcenter', 'border': 1,
            'bg_color': '#112233', 'font_color': '#FFFFFF', 'text_wrap': True
        })
        firma_label_format = workbook.add_format({
            'bold': True, 'align': 'right', 'valign': 'vcenter',
            'bg_color': '#1E293B', 'font_color': '#94A3B8',
            'font_size': 9, 'border': 0
        })
        firma_value_format = workbook.add_format({
            'bold': True, 'align': 'left', 'valign': 'vcenter',
            'bg_color': '#0F172A', 'font_color': '#F1F5F9',
            'font_size': 10, 'border': 0
        })
        firma_title_format = workbook.add_format({
            'bold': True, 'align': 'center', 'valign': 'vcenter',
            'bg_color': '#1E3A5F', 'font_color': '#FFFFFF',
            'font_size': 14, 'border': 0
        })

        # ----- RİSK SEVİYESİ RENK FORMATLARI -----
        fmt_red    = workbook.add_format({'bg_color': '#FF0000', 'font_color': '#FFFFFF', 'bold': True, 'align': 'center', 'valign': 'vcenter', 'border': 1})
        fmt_yellow = workbook.add_format({'bg_color': '#FFFF00', 'font_color': '#000000', 'bold': True, 'align': 'center', 'valign': 'vcenter', 'border': 1})
        fmt_blue   = workbook.add_format({'bg_color': '#0070C0', 'font_color': '#FFFFFF', 'bold': True, 'align': 'center', 'valign': 'vcenter', 'border': 1})
        fmt_green  = workbook.add_format({'bg_color': '#00B050', 'font_color': '#000000', 'bold': True, 'align': 'center', 'valign': 'vcenter', 'border': 1})

        def get_color_format(score):
            if score is None: return base_format
            try:
                s = float(score)
            except:
                return base_format
            if s > 400: return fmt_red
            elif s > 200: return fmt_yellow
            elif s > 70: return fmt_blue
            elif s > 20: return fmt_green
            else: return base_format

        # -----------------------------
        # FİRMA BİLGİLERİ BLOĞU (Üst kısım)
        # -----------------------------
        TOTAL_COLS = 18

        # Satır 0: Başlık (tam genişlik)
        worksheet.merge_range(0, 0, 0, TOTAL_COLS - 1, 'RİSK DEĞERLENDİRME RAPORU', firma_title_format)
        worksheet.set_row(0, 30)

        # Satırlar 1..N: Firma bilgileri (2 sütun: etiket + değer, geri kalan boş)
        for i, (label, value) in enumerate(firma_info):
            row_i = i + 1
            worksheet.set_row(row_i, 18)
            worksheet.write(row_i, 0, label, firma_label_format)
            worksheet.merge_range(row_i, 1, row_i, TOTAL_COLS - 1, value, firma_value_format)

        # Boş ayırıcı satır (DATA_HEADER_ROW - 1)
        worksheet.set_row(FIRMA_ROWS, 6)

        # -----------------------------
        # BAŞLIK (HEADER) BİRLEŞTİRMELERİ
        # -----------------------------
        H0 = DATA_HEADER_ROW
        H1 = DATA_HEADER_ROW + 1

        vertical_merges = [
            (0, 'Sıra No'),
            (1, 'Yapılan İş / Faaliyet / Ekipman / Ortam'),
            (2, 'Tehlike'),
            (3, 'Risk / Olası Etki'),
            (4, 'Mevcutta Alınan Önlemler'),
            (10, 'Açıklama / Düzeltici ve Önleyici Tedbirler'),
            (16, 'Sorumlu'),
            (17, 'Termin Süresi')
        ]
        for col_idx_num, title in vertical_merges:
            worksheet.merge_range(H0, col_idx_num, H1, col_idx_num, title, header_format)

        worksheet.merge_range(H0, 5, H0, 9,  'İLK RİSK DEĞERLERİ',  header_format)
        worksheet.merge_range(H0, 11, H0, 15, 'SON RİSK DEĞERLERİ', header_format)

        sub_headers = ['Olasılık', 'Frekans', 'Şiddet', 'Risk Puanı (R)', 'Risk Düzeyi']
        for idx, sub_title in enumerate(sub_headers):
            worksheet.write(H1, 5  + idx, sub_title, header_format)
            worksheet.write(H1, 11 + idx, sub_title, header_format)

        # -----------------------------
        # SÜTUN GENİŞLİKLERİ
        # -----------------------------
        col_idx = {name: i for i, name in enumerate(columns)}

        worksheet.set_column(0,  0,  8,  base_format)
        worksheet.set_column(1,  1,  30, wrap_format)
        worksheet.set_column(2,  2,  40, wrap_format)
        worksheet.set_column(3,  3,  50, wrap_format)
        worksheet.set_column(4,  4,  25, wrap_format)
        worksheet.set_column(5,  9,  12, base_format)
        worksheet.set_column(10, 10, 60, wrap_format)
        worksheet.set_column(11, 15, 12, base_format)
        worksheet.set_column(16, 17, 20, base_format)

        # Verileri renklendirerek yaz
        for row_num, row_data in enumerate(rows):
            excel_row = row_num + DATA_START_ROW
            for col_num, cell_value in enumerate(row_data):
                current_format = base_format
                if col_num in [1, 2, 3, 4, 10]:
                    current_format = wrap_format
                if col_num in [8, 9]:
                    score = row_data[col_idx['İlk Risk - Risk Puanı']]
                    current_format = get_color_format(score)
                if col_num in [14, 15]:
                    score = row_data[col_idx['Son Risk - Risk Puanı']]
                    current_format = get_color_format(score)
                value_to_write = "" if pd.isna(cell_value) else cell_value
                worksheet.write(excel_row, col_num, value_to_write, current_format)

    output.seek(0)
    return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True, download_name='risk_degerlendirmesi_formati.xlsx')


# --- WORD EXPORT ---

@app.route('/api/companies/<int:company_id>/export/docx', methods=['GET'])
def export_company_docx(company_id):
    """Firma için kapsamlı Word belgesi oluştur."""
    from docx import Document as DocxDocument
    from docx.shared import RGBColor, Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from datetime import date
    import copy

    company = Company.query.get_or_404(company_id)

    # --- Yardımcı Fonksiyonlar ---
    def fmt_date(d):
        if not d: return '—'
        if isinstance(d, str):
            try:
                parts = d.split('-')
                return f"{parts[2]}.{parts[1]}.{parts[0]}"
            except: return d
        return d.strftime('%d.%m.%Y')

    def cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def cell_valign(cell, align='center'):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), align)
        tcPr.append(vAlign)

    def set_col_width(table, col_idx, width_cm):
        for row in table.rows:
            row.cells[col_idx].width = Cm(width_cm)

    def header_para(doc, text, level=1):
        """Renkli bölüm başlığı paragrafı ekle"""
        colors = {1: '1E3A5F', 2: '1E40AF', 3: '374151'}
        sizes  = {1: 16, 2: 13, 3: 11}
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(sizes.get(level, 11))
        run.font.color.rgb = RGBColor.from_string(colors.get(level, '1E3A5F'))
        if level == 1:
            # Alt çizgi ekle
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '4')
            bottom.set(qn('w:color'), '6366F1')
            pBdr.append(bottom)
            pPr.append(pBdr)
        return p

    def add_info_table(doc, rows_data, col_widths=(5, 12)):
        """Etiket-Değer çift sütunlu bilgi tablosu"""
        table = doc.add_table(rows=len(rows_data), cols=2)
        table.style = 'Table Grid'
        total = sum(col_widths)
        for i, (label, value) in enumerate(rows_data):
            row = table.rows[i]
            lc, vc = row.cells[0], row.cells[1]
            lc.width = Cm(col_widths[0])
            vc.width = Cm(col_widths[1])
            cell_bg(lc, 'EFF6FF')
            cell_bg(vc, 'FFFFFF')
            cell_valign(lc)
            cell_valign(vc)
            lp = lc.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            lr = lp.add_run(str(label))
            lr.font.bold = True
            lr.font.size = Pt(9)
            lr.font.color.rgb = RGBColor.from_string('1E40AF')
            vp = vc.paragraphs[0]
            vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            vr = vp.add_run(str(value) if value else '—')
            vr.font.size = Pt(9)
            vr.font.color.rgb = RGBColor.from_string('111827')
        doc.add_paragraph()

    def add_data_table(doc, headers, rows, header_color='1E3A5F'):
        """Başlıklı veri tablosu"""
        if not rows:
            p = doc.add_paragraph('Kayıt bulunamadı.')
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor.from_string('6B7280')
            doc.add_paragraph()
            return
        ncols = len(headers)
        table = doc.add_table(rows=1 + len(rows), cols=ncols)
        table.style = 'Table Grid'
        # Header satırı
        hrow = table.rows[0]
        for j, h in enumerate(headers):
            c = hrow.cells[j]
            cell_bg(c, header_color)
            cell_valign(c)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(h))
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string('FFFFFF')
        # Veri satırları
        for i, row_data in enumerate(rows):
            drow = table.rows[i + 1]
            bg = 'F8FAFC' if i % 2 == 0 else 'FFFFFF'
            for j, val in enumerate(row_data):
                c = drow.cells[j]
                cell_bg(c, bg)
                cell_valign(c)
                p = c.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(str(val) if val is not None else '—')
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor.from_string('1F2937')
        doc.add_paragraph()

    # --- Veri Toplama ---
    tehlike = company.nace_code.danger_class.name if company.nace_code and company.nace_code.danger_class else '—'
    nace = f"{company.nace_code.code} — {company.nace_code.description}" if company.nace_code else '—'
    city_name = company.city.name if company.city else '—'
    dist_name = company.district.name if company.district else '—'
    full_address = f"{company.address or ''}, {dist_name} / {city_name}".strip(', ')

    personnel_list = Personnel.query.filter_by(company_id=company_id).order_by(Personnel.last_name).all()
    professionals  = CompanyProfessional.query.filter_by(company_id=company_id).all()
    assignments    = Assignment.query.filter_by(company_id=company_id).all()
    risks          = RiskAssessment.query.filter_by(company_id=company_id).order_by(RiskAssessment.id).all()

    # Validity hesaplama
    validity_years = {'Çok Tehlikeli': 1, 'Tehlikeli': 3, 'Az Tehlikeli': 5}
    health_validity_years = {'Çok Tehlikeli': 1, 'Tehlikeli': 3, 'Az Tehlikeli': 5}

    def get_health_status(p):
        if not p.health_exam_date: return 'Muayene Yok'
        if not p.health_validity_date: return 'Muayene Yok'
        today = date.today()
        diff = (p.health_validity_date - today).days
        if diff < 0: return 'Süresi Dolmuş'
        if diff <= 90: return 'Yaklaşan'
        return 'Geçerli'

    def get_training_status(p):
        vd = p.training_validity_date
        if not vd: return 'Eğitim Yok'
        today = date.today()
        diff = (vd - today).days
        if diff < 0: return 'Süresi Dolmuş'
        if diff <= 90: return 'Yaklaşan'
        return 'Geçerli'

    # --- Word Belgesi ---
    doc = DocxDocument()

    # Sayfa kenar boşlukları
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # ---- KAPAK SAYFASI ----
    doc.add_paragraph()
    doc.add_paragraph()

    # Büyük gradient-benzeri başlık kutusu
    cover_table = doc.add_table(rows=1, cols=1)
    cover_table.style = 'Table Grid'
    cover_cell = cover_table.rows[0].cells[0]
    cover_cell.width = Cm(17)
    cell_bg(cover_cell, '1E3A5F')
    cover_cell._tc.get_or_add_tcPr()
    # İç padding
    p1 = cover_cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(20)
    p1.paragraph_format.space_after  = Pt(4)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run('FİRMA BİLGİ KARTI')
    r1.font.bold = True
    r1.font.size = Pt(22)
    r1.font.color.rgb = RGBColor.from_string('FFFFFF')
    r1.font.name = 'Calibri'

    p2 = cover_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(company.official_title)
    r2.font.size = Pt(14)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor.from_string('93C5FD')

    p3 = cover_cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(20)
    r3 = p3.add_run(f"Oluşturulma Tarihi: {date.today().strftime('%d.%m.%Y')}  |  Tehlike Sınıfı: {tehlike}")
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor.from_string('BFDBFE')

    doc.add_paragraph()
    doc.add_paragraph()

    # İkinci renkli bant — özet istatistikler
    stats_table = doc.add_table(rows=1, cols=3)
    stats_table.style = 'Table Grid'
    stats_data = [
        ('Çalışan Sayısı', str(len(personnel_list)), '2563EB'),
        ('Risk Kaydı', str(len(risks)), '7C3AED'),
        ('Görevlendirme', str(len(assignments)), '059669'),
    ]
    for j, (label, val, color) in enumerate(stats_data):
        c = stats_table.rows[0].cells[j]
        cell_bg(c, color)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(f"{val}\n{label}")
        r.font.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor.from_string('FFFFFF')

    doc.add_page_break()

    # ---- BÖLÜM 1: FİRMA GENEL BİLGİLERİ ----
    header_para(doc, '1. Firma Genel Bilgileri', level=1)
    add_info_table(doc, [
        ('Firma Unvanı',          company.official_title),
        ('SGK Sicil No',          company.sgk_no),
        ('NACE Kodu',             nace),
        ('Tehlike Sınıfı',        tehlike),
        ('Adres',                 full_address),
        ('Çalışan Sayısı',        str(len(personnel_list))),
        ('İşveren / Vekili',      company.employer_name or '—'),
        ('İşveren Görevi',        company.employer_role or '—'),
        ('İşveren Telefonu',      company.employer_phone or '—'),
        ('İşveren E-posta',       company.employer_email or '—'),
        ('Yetkili Personel',      company.authorized_person or '—'),
        ('Yetkili Görevi',        company.authorized_role or '—'),
        ('Yetkili Telefonu',      company.authorized_phone or '—'),
        ('Yetkili E-posta',       company.authorized_email or '—'),
    ])

    # ---- BÖLÜM 2: İSG PROFESYONELLERİ ----
    header_para(doc, '2. İSG Profesyonelleri', level=1)
    role_labels = {'IGU': 'İş Güvenliği Uzmanı', 'IYH': 'İşyeri Hekimi', 'DSP': 'Diğer Sağlık Personeli'}
    prof_rows = [[
        p.professional_name,
        role_labels.get(p.role.name, p.role.name) if p.role else '—',
        p.certificate_class or '—',
        p.certificate_number or '—',
        fmt_date(p.contract_start_date),
        'Aktif' if p.is_active else 'Pasif'
    ] for p in professionals]
    add_data_table(doc, ['Ad Soyad', 'Görev', 'Sertifika Sınıfı', 'Belge No', 'Sözleşme Tarihi', 'Durum'], prof_rows, '1E3A5F')

    # ---- BÖLÜM 3: PERSONEL LİSTESİ ----
    header_para(doc, '3. Personel Listesi', level=1)
    pers_rows = [[
        f"{p.first_name} {p.last_name}",
        p.tc_no or '—',
        p.job_title or '—',
        fmt_date(p.hire_date),
        p.blood_type or '—',
        p.phone_number or '—',
    ] for p in personnel_list]
    add_data_table(doc, ['Ad Soyad', 'TC Kimlik No', 'Görevi', 'İşe Giriş', 'Kan Grubu', 'Telefon'], pers_rows, '1E40AF')

    # ---- BÖLÜM 4: İSG EĞİTİMLERİ ----
    header_para(doc, '4. İSG Eğitimleri', level=1)
    train_rows = [[
        f"{p.first_name} {p.last_name}",
        p.tc_no or '—',
        p.job_title or '—',
        fmt_date(p.training_date_1),
        fmt_date(p.training_date_2),
        fmt_date(p.training_validity_date),
        get_training_status(p),
    ] for p in personnel_list]
    add_data_table(doc,
        ['Ad Soyad', 'TC Kimlik No', 'Görevi', 'Eğitim 1', 'Eğitim 2', 'Geçerlilik', 'Durum'],
        train_rows, '065F46')

    # ---- BÖLÜM 5: SAĞLIK GÖZETİMİ ----
    header_para(doc, '5. Sağlık Gözetimi', level=1)
    health_rows = [[
        f"{p.first_name} {p.last_name}",
        p.tc_no or '—',
        p.job_title or '—',
        fmt_date(p.health_exam_date),
        fmt_date(p.health_validity_date),
        get_health_status(p),
    ] for p in personnel_list]
    add_data_table(doc,
        ['Ad Soyad', 'TC Kimlik No', 'Görevi', 'Muayene Tarihi', 'Geçerlilik', 'Durum'],
        health_rows, '92400E')

    # ---- BÖLÜM 6: RİSK DEĞERLENDİRMESİ ----
    header_para(doc, '6. Risk Değerlendirmesi', level=1)

    # Özet istatistik
    risk_counts = {'Tolerans Gösterilemez Risk': 0, 'Esaslı Risk': 0, 'Önemli Risk': 0, 'Olası Risk': 0, 'Önemsiz Risk': 0}
    for r in risks:
        lvl = r.final_risk_level or r.initial_risk_level or '—'
        if lvl in risk_counts: risk_counts[lvl] += 1

    summary_rows = [[lvl, str(cnt)] for lvl, cnt in risk_counts.items()]
    add_info_table(doc, [('Toplam Risk Kaydı', str(len(risks)))] + [(k, v) for k, v in risk_counts.items()], col_widths=(6, 3))

    risk_rows = [[
        str(idx + 1),
        (r.activity or '')[:60],
        (r.hazard or '')[:50],
        (r.risk_effect or '')[:40],
        str(r.initial_risk_score or '—'),
        r.initial_risk_level or '—',
        str(r.final_risk_score or '—'),
        r.final_risk_level or '—',
        r.responsible_person or '—',
        r.deadline or '—',
    ] for idx, r in enumerate(risks)]
    add_data_table(doc,
        ['#', 'Faaliyet', 'Tehlike', 'Risk / Etki', 'İlk R', 'İlk Düzey', 'Son R', 'Son Düzey', 'Sorumlu', 'Termin'],
        risk_rows, '7C3AED')

    # ---- BÖLÜM 7: GÖREVLENDİRMELER ----
    header_para(doc, '7. Görevlendirmeler', level=1)
    assign_rows = []
    for a in assignments:
        name = '—'
        if a.personnel_id:
            p = next((x for x in personnel_list if x.id == a.personnel_id), None)
            if p: name = f"{p.first_name} {p.last_name}"
        elif a.professional_id:
            pro = next((x for x in professionals if x.id == a.professional_id), None)
            if pro: name = pro.professional_name
        assign_rows.append([
            name,
            a.assignment_type.value if a.assignment_type else '—',
            fmt_date(a.assignment_date) if hasattr(a, 'assignment_date') else '—',
        ])
    add_data_table(doc, ['Ad Soyad', 'Görev Türü', 'Tarih'], assign_rows, '1E3A5F')

    # ---- FOOTER: Sayfa Numarası ----
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.add_run(f"{company.official_title}  |  Oluşturulma: {date.today().strftime('%d.%m.%Y')}").font.size = Pt(8)

    # BytesIO'ya yaz
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    safe_name = company.official_title.replace(' ', '_').replace('/', '-')[:40]
    return send_file(out,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True,
                     download_name=f'firma_bilgi_karti_{safe_name}.docx')


# --- RISK WIZARD ROUTES ---
import json as json_module

@app.route('/api/risk-library', methods=['GET'])
def get_risk_library():
    lib_path = os.path.join(basedir, 'risk_library.json')
    with open(lib_path, 'r', encoding='utf-8') as f:
        library = json_module.load(f)
    # Return only category names and icons for checklist
    categories = []
    for cat in library:
        categories.append({
            'category': cat['category'],
            'icon': cat.get('icon', ''),
            'itemCount': len(cat['items'])
        })
    return jsonify(categories), 200

@app.route('/api/risk-library/<path:category_name>', methods=['GET'])
def get_risk_library_category(category_name):
    lib_path = os.path.join(basedir, 'risk_library.json')
    with open(lib_path, 'r', encoding='utf-8') as f:
        library = json_module.load(f)
    for cat in library:
        if cat['category'] == category_name:
            return jsonify(cat['items']), 200
    return jsonify([]), 404

def decrease_value(val):
    options = [0.2, 0.5, 1, 3, 6, 10]
    for i in range(len(options)):
        if val <= options[i]:
            return options[max(0, i - 1)]
    return 6

@app.route('/api/companies/<int:company_id>/risks/generate', methods=['POST'])
def generate_risks_from_wizard(company_id):
    Company.query.get_or_404(company_id)
    data = request.json
    selected_categories = data.get('categories', [])
    
    if not selected_categories:
        return jsonify({'error': 'En az bir kategori seçilmelidir.'}), 400
    
    lib_path = os.path.join(basedir, 'risk_library.json')
    with open(lib_path, 'r', encoding='utf-8') as f:
        library = json_module.load(f)
    
    count = 0
    try:
        for cat in library:
            if cat['category'] in selected_categories:
                for item in cat['items']:
                    init_p = float(item.get('default_O', 0))
                    init_s = float(item.get('default_S', 0))
                    init_f = float(item.get('default_F', 0))
                    
                    risk = RiskAssessment(
                        company_id=company_id,
                        activity=item.get('activity', ''),
                        hazard=item.get('hazard', ''),
                        risk_effect=item.get('risk', ''),
                        corrective_measures=item.get('measures', ''),
                        existing_measures='',
                        initial_probability=init_p,
                        initial_severity=init_s,
                        initial_frequency=init_f,
                        final_probability=decrease_value(init_p),
                        final_severity=init_s, # Sabit
                        final_frequency=decrease_value(init_f),
                        responsible_person='',
                        deadline='',
                    )
                    risk.calculate_risks()
                    db.session.add(risk)
                    count += 1
        db.session.commit()
        return jsonify({'message': f'{count} risk kaydı başarıyla oluşturuldu.', 'count': count}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")
OPENROUTER_MODEL = os.getenv("OPENROUTER_MODEL", "google/gemini-3.1-pro-preview")

if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

def call_openrouter(system_instruction, prompt_text):
    """Calls OpenRouter API with the given system instruction and prompt."""
    if not OPENROUTER_API_KEY:
        return None
    
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "HTTP-Referer": os.environ.get("FRONTEND_URL", "http://localhost:5173"),
        "X-Title": "ISG Tracking System",
        "Content-Type": "application/json"
    }
    
    data = {
        "model": OPENROUTER_MODEL,
        "messages": [
            {"role": "system", "content": system_instruction},
            {"role": "user", "content": prompt_text}
        ],
        "temperature": 0.7,
        "max_tokens": 8192
    }
    
    try:
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=120
        )
        response.raise_for_status()
        result = response.json()
        return result['choices'][0]['message']['content'].strip()
    except Exception as e:
        print(f"OpenRouter Error: {e}")
        return None

@app.route('/api/generate-regulation-checklist', methods=['POST'])
def generate_regulation_checklist():
    """Verilen yönetmelik için profesyonel denetim kontrol listesi üretir."""
    if not GEMINI_API_KEY:
        return jsonify({'error': 'Sunucuda API Key bulunamadı.'}), 500

    data = request.get_json()
    title    = data.get('title', '')
    category = data.get('category', '')
    content  = data.get('content', '')

    system_instruction = """Sen Türkiye'deki İSG mevzuatı ve teknik standartlar konusunda 20+ yıl deneyimli, lisanslı A sınıfı İSG uzmanı ve baş denetçisin.
Eğitim aldığın ve yürürlükteki tüm Türk İSG yönetmelikleri, Çalışma ve Sosyal Güvenlik Bakanlığı tebliğleri ve ilgili teknik standartlar hakkında eksiksiz bilgiye sahipsin.

━━━ GÖREV TANIMI ━━━
Verilen yönetmelik için SAHADAKİ denetçinin birebir kullanabileceği, teknik derinlikte, madde numarası referanslı, SAYISAL EŞİK DEĞERLERİ İÇEREN eksiksiz bir denetim kontrol listesi üret.

━━━ AŞAMA 1 — YÖNETMELİĞİ ZİHİNDE ANALİZ ET ━━━
Kontrol listesini oluşturmadan önce aşağıdaki soruları YANITLA (bu yanıtları JSON'a YAZMA, sadece düşün):
a) Bu yönetmelik hangi spesifik konuları/varlıkları düzenliyor? (örn: ekranlı araçlarda → ekran, klavye, masa, sandalye, mola, göz muayenesi, aydınlatma, sıcaklık, nem, gürültü)
b) Yönetmeliğin tüm bölüm/madde başlıkları neler? Her birini listele.
c) Hangi sayısal eşik değerleri var? (örn: max. lüks, min. cm, periyot, sıcaklık aralığı)
d) Hangi belge/kayıt tutma yükümlülükleri var?
e) Hangi periyodik kontrol/test/muayene zorunlulukları var?
f) Hangi eğitim ve yetkilendirme gereklilikleri var?
Şimdi bu analize dayanarak JSON'u oluştur.

━━━ YAZIM KURALLARI ━━━
• Eylem/durum ifadesiyle BİTMELİ: "yapılmıştır / mevcuttur / sağlanmıştır / belgelenmiştir / uygulanmaktadır / doğrulanmıştır / tutulmaktadır / kontrol edilmiştir"
• DOĞRU: "Çalışma yüzeyi yüksekliği 72-78 cm aralığında ayarlanabilir nitelikte olup ölçüm ile doğrulanmıştır."
• YANLIŞ: "Masa yüksekliği uygun mu?"
• Denetçi sahada tek bakışta "Uygun / Uygun Değil / Uygulanmıyor" diyebilecek kadar somut olmalı

━━━ SAYISAL DEĞER ZORUNLULUĞU ━━━
Her maddede mümkünse şunlardan birini yaz:
• Boyut/mesafe: "en az 120 cm", "min. 72 cm", "maks. 10 m²/kişi"
• Sıcaklık/nem: "20-26°C", "%40-70 bağıl nem"
• Aydınlatma: "en az 300 lüks", "500 lüks iş yüzeyinde", "kamaşma indeksi UGR<19"
• Gürültü: "85 dB(A) maruziyet sınırı", "Lep,d maks. 87 dB(A)"
• Periyot/süre: "yılda 1 kez", "2 saatte bir 15 dakika", "her 4 yılda"
• Kapasite: "6 kg KKÖ", "1 adet / 200 m²", "min. 34A-233B-C"
• Elektrik/basınç: "230 V ±10%", "IP44", "6 bar"
• Tarih: "son 12 ay içinde", "geçerlilik süresi dolmamış"

━━━ KAPSAM KURALLARI ━━━
• Her bölüm 8-12 madde — AZ MADDE KESİNLİKLE KABUL EDİLEMEZ
• Minimum 8, maksimum 15 bölüm — her alt konu ayrı bölüm
• Yönetmeliğe ÖZGÜ tüm varlıkları kapsayan bölümler yap:
  - Ekranlı Araçlar → Ekran, Klavye/Mouse, Masa/Çalışma yüzeyi, Sandalye/Oturma, Aydınlatma, Isı/Nem/Hava, Gürültü, Yazılım/Arayüz, Mola programı, Göz muayenesi, Hamile/Risk grubu
  - Yangından Korunma → Kaçış yolları, Sprinkler, Yangın kapıları, Söndürücüler, Tahliye planı, Bina sınıflandırma, Elektrik tesisatı, Acil aydınlatma, Yangın kompartımanı, Tatbikat
  - Kimyasal Maddeler → Depolama, Etiketleme, GBF, KKD, Maruziyet ölçümü, Atık yönetimi, Acil durum, Sağlık gözetimi
  (diğer yönetmelikler için de benzer özel analiz yap)
• Madde numaraları kesin olarak doğru olmalı — bilmiyorsan "—" yaz
• TS EN / ISO / IEC / NFPA / TSE standartlarına somut atıf
• Belge/kayıt yükümlülükleri, periyodik kontroller, eğitim ayrı maddeler
• İdari para cezası riski taşıyan maddeler mutlaka dahil
• Türkçe, teknik terminoloji
• SADECE JSON — başka hiçbir metin ekleme

━━━ JSON FORMAT ━━━
[
  {
    "section": "Bölüm Başlığı",
    "icon": "AŞAĞIDAN_BİR_İKON",
    "items": [
      {
        "madde": "Md. X/Y veya —",
        "kontrol": "Sayısal eşik içeren, eylem/durum ifadesiyle biten somut tespit cümlesi",
        "referans": "TS EN XXXX / ISO YYYY / ilgili teknik gereklilik ve limit değeri",
        "risk_seviyesi": "yüksek|orta|düşük"
      }
    ]
  }
]

İKON LİSTESİ (SADECE bunlardan seç):
description, gavel, health_and_safety, engineering, warning, fire_extinguisher, local_fire_department,
medical_services, school, assignment, fact_check, verified_user, security, policy, report,
build, construction, electrical_services, bolt, water, air, thermostat, speed, science,
inventory_2, warehouse, local_shipping, directions_car, business_center, badge, groups,
monitor_heart, emergency, accessible, visibility, lock, key, admin_panel_settings,
folder, article, checklist, task_alt, done_all, rule, list_alt, format_list_bulleted,
settings, tune, handyman, plumbing, architecture, foundation, domain, apartment,
eco, recycling, compost, public, grass, agriculture, landscape, wb_sunny,
electric_bolt, cable, power, battery_charging_full, sensors, router, devices,
calculate, straighten, scale, biotech, microscope, lab_research, hub"""

    # İçerik çok uzunsa kırp
    if content and len(content) > 4000:
        content = content[:4000] + '\n[...metin kısaltıldı]'

    content_section = ('Yönetmelik Metni/Özeti:\n' + content) if content else ''
    prompt = f"""YÖNETMELİK: {title}
KATEGORİ: {category}
{content_section}

GÖREV: Bu yönetmeliğin tüm konularını (madde madde) zihinsel olarak tara. Yönetmeliğin adından ve bilgi birikiminden yararlanarak:

▶ TEMEL METİN ANALİZİ:
• Bu yönetmeliğin düzenlediği TÜM varlıkları/ekipmanları/ortamları belirle
  (örn. ekranlı araçlarda: monitör, klavye, fare, masa, sandalye, ayak desteği, aydınlatma, panjur/perde, klima, gürültü kaynakları, yazılım arayüzü)
• Her bölüm için yönetmeliğe ÖZGÜ maddeler yaz — genel/jenerik maddeler kabul edilemez
• BU yönetmeliğin getirdiği spesifik yükümlülükleri sorgula

▶ EK/CETVEL ANALİZİ (ÇOK ÖNEMLİ):
Bu yönetmeliğin eklerini/cetvellerini de bilgi birikiminden yararlanarak analiz et:
• Minimum gereklilik tabloları (boyut, kapasite, mesafe cetvelleri)
• Teknik şartname ekleri (malzeme sınıfları, test metotları)
• Sınıflandırma/kategorilendirme tabloları (risk grupları, yapı sınıfları, tehlike kategorileri)
• Formlar ve örnek belgeler (kayıt formları, bildirim şablonları)
• Standart değerler listesi (maruziyet sınır değerleri, eylem değerleri)
Eklerdeki her cetvel/tablo için ayrı bölüm veya maddeler oluştur.

▶ ZORUNLU KAPSAM (hepsini dahil et):
① Fiziksel/teknik gereklilikler — boyut, ağırlık, kapasite, malzeme sınıfı, dayanım süreleri
② Çevre koşulları — sıcaklık, nem, aydınlatma (lüks), gürültü (dB), titreşim, hava kalitesi
③ Ergonomi ve çalışma düzeni — bu yönetmeliğe özgü pozisyon/donanım gereklilikleri
④ Periyodik muayene/test/ölçüm — ne, ne zaman, kim yapar, hangi belgeyle kanıtlanır
⑤ Sağlık gözetimi — muayene türü, periyodu, özel gruplar (hamile, genç, engelli)
⑥ Eğitim ve yetkilendirme — içerik, süre (saat), periyot, sertifika türü
⑦ Belge ve kayıt tutma — tür, saklama süresi, format, imzalayanlar
⑧ İşaret/etiket/uyarı — renk kodu, boyut, dil, yerleştirme yeri
⑨ Acil durum prosedürleri — eylem planı, tatbikat sıklığı, sorumlular
⑩ Yönetimsel/idari — komisyon, görevlendirme, bildirim, idari para cezası riski

ÇIKTI: Sadece JSON — başka hiçbir şey yazma."""

    def _parse_checklist_json(text: str):
        """JSON parse + bracket extraction + normalize. sections listesi döndürür, başarısız olursa None."""
        text = text.strip()
        # Direkt parse
        for attempt_text in [text]:
            try:
                parsed = json.loads(attempt_text)
                break
            except Exception:
                parsed = None

        if parsed is None:
            # Bracket extraction
            arr_start = text.find('[')
            arr_end   = text.rfind(']')
            if arr_start != -1 and arr_end != -1 and arr_start < arr_end:
                candidate = text[arr_start:arr_end + 1]
                candidate = re.sub(r',\s*([}\]])', r'\1', candidate)  # trailing commas
                # Kırpılmış JSON'u kapat
                open_b = candidate.count('{') - candidate.count('}')
                open_s = candidate.count('[') - candidate.count(']')
                candidate += '}' * max(0, open_b) + ']' * max(0, open_s)
                try:
                    parsed = json.loads(candidate)
                except Exception:
                    pass

        if parsed is None:
            # Object extraction
            obj_start = text.find('{')
            obj_end   = text.rfind('}')
            if obj_start != -1 and obj_end != -1:
                candidate = text[obj_start:obj_end + 1]
                candidate = re.sub(r',\s*([}\]])', r'\1', candidate)
                try:
                    parsed = json.loads(candidate)
                except Exception:
                    pass

        if parsed is None:
            return None

        # Normalize: list veya {"sections":[...]} kabul et
        if isinstance(parsed, list):
            return parsed
        if isinstance(parsed, dict):
            for key in ('sections', 'checklist', 'items', 'data'):
                if key in parsed and isinstance(parsed[key], list):
                    return parsed[key]
            vals = [v for v in parsed.values() if isinstance(v, list)]
            if vals:
                return vals[0]
        return None

    model = genai.GenerativeModel('gemini-2.5-flash', system_instruction=system_instruction)
    gen_config = genai.types.GenerationConfig(
        temperature=0.2,          # Daha düşük → daha tutarlı, daha az halüsinasyon
        max_output_tokens=14000,  # Kapsamlı yönetmelikler için yeterli
        response_mime_type="application/json"
    )

    MAX_RETRIES = 3
    last_error  = None

    for attempt in range(MAX_RETRIES):
        try:
            response = model.generate_content(prompt, generation_config=gen_config)
            text     = response.text.strip() if response.text else ''
            app.logger.info(
                f"[attempt {attempt+1}] Regulation checklist ham çıktı "
                f"({len(text)} kar, ilk 150): {text[:150]}"
            )

            sections = _parse_checklist_json(text)
            if sections and len(sections) > 0:
                # Her section'ın zorunlu alanları var mı kontrol et
                valid = [
                    s for s in sections
                    if isinstance(s, dict) and s.get('section') and isinstance(s.get('items'), list)
                ]
                if valid:
                    app.logger.info(
                        f"Regulation checklist üretildi: {title[:50]} → "
                        f"{len(valid)} bölüm (attempt {attempt+1})"
                    )
                    return jsonify({'sections': valid})

            app.logger.warning(f"[attempt {attempt+1}] Geçerli section bulunamadı, tekrar deneniyor…")
            last_error = 'Model geçerli bölüm içeriği üretemedi'

        except Exception as e:
            last_error = str(e)
            app.logger.error(f"[attempt {attempt+1}] Regulation checklist hatası: {e}")
            if attempt < MAX_RETRIES - 1:
                time.sleep(2 * (attempt + 1))  # 2s, 4s backoff

    app.logger.error(f"Regulation checklist {MAX_RETRIES} denemede başarısız: {last_error}")
    return jsonify({'error': f'Kontrol listesi oluşturulamadı: {last_error}'}), 500


@app.route('/api/generate-checklist', methods=['POST'])
def generate_checklist_ai():
    if not GEMINI_API_KEY:
        return jsonify({'error': 'Sunucuda API Key bulunamadı.'}), 500

    data = request.json
    prompt_text = data.get('prompt', '')
    scope = data.get('scope', 'general') # default: general
    
    if not prompt_text:
        return jsonify({'error': 'Faaliyet/Ortam bilgisi boş olamaz.'}), 400

    system_instruction = """Sen son derece profesyonel, kıdemli bir İSG (İş Sağlığı ve Güvenliği) Baş Mühendisisin. Kullanıcının girdiği tesis/işletme tipini (veya makineyi) analiz et.
Bu tesis için İSG standartlarına uygun, bölüm bölüm (veya süreç süreç) ayrılmış kapsamlı bir "Saha Denetim ve Risk Planlama Rehberi" oluştur.
Mümkün olan tüm tehlikeleri listele ve mantıklı kategorilere böl.

Çıktı FORMATI: YALNIZCA geçerli bir JSON array olmalıdır, başka hiçbir markdown, açıklama veya kelime içermemelidir.
Her bölüm için aşağıdaki JSON yapısına sadık kalmalısın:
[
  {
    "bolum_adi": "Örn: Hammadde Kabul ve Depolama",
    "bakilacak_tehlikeler": ["Forklift trafiği", "Yüksekten düşme", "Malzeme devrilmesi"],
    "tipik_tedbir_mantigi": "İstif standartları (TS EN ...), bariyerleme uygulamaları ve yaya-araç yolu ayrımı."
  },
  {
    "bolum_adi": "Örn: Kesim / Üretim Alanı",
    "bakilacak_tehlikeler": ["Uzuv sıkışması", "Gürültü (85 dB üstü)"],
    "tipik_tedbir_mantigi": "Makine koruyucuları (TS EN ISO 12100) ve Kişisel Koruyucu Donanım kullanımı."
  }
]
"""

    if scope in ['specific', 'machine']:
        system_instruction += """
DİKKAT KATI KURAL: Kullanıcı senden sadece çok spesifik bir konuyu (veya makineyi) analiz etmeni istiyor. KESİNLİKLE KAPSAM DIŞINA ÇIKMA. Tesisin diğer genel risklerini (eğer kullanıcı özellikle yazmadıysa) ASLA listeye ekleme. Sadece ve sadece kullanıcının girdiği metne %100 sadık kalarak, o dar kapsamdaki mikro tehlikeleri derinlemesine analiz et.
"""
    try:
        # 1. Try OpenRouter first if available
        if OPENROUTER_API_KEY:
            text = call_openrouter(system_instruction, prompt_text)
            if text:
                # Robust JSON Array Extraction
                start_idx = text.find('[')
                end_idx = text.rfind(']')
                if start_idx != -1 and end_idx != -1 and start_idx <= end_idx:
                    text = text[start_idx:end_idx+1]
                    # LLM JSON Hatalarını temizle
                    text = re.sub(r',\s*}', '}', text)
                    text = re.sub(r',\s*]', ']', text)
                    try:
                        return jsonify(json.loads(text)), 200
                    except:
                        pass # Fallback to Gemini SDK below

        # 2. Fallback to Gemini SDK
        if not GEMINI_API_KEY:
            return jsonify({'error': 'Sunucuda API Key bulunamadı.'}), 500

        model = genai.GenerativeModel('gemini-2.5-flash', system_instruction=system_instruction)
        max_retries = 3
        checklist = []
        
        for attempt in range(max_retries):
            try:
                response = model.generate_content(
                    prompt_text,
                    generation_config=genai.types.GenerationConfig(
                        temperature=0.7,
                        max_output_tokens=8192,
                        response_mime_type="application/json"
                    )
                )
                text = response.text.strip()
                
                # Robust JSON Array Extraction
                start_idx = text.find('[')
                end_idx = text.rfind(']')
                if start_idx != -1 and end_idx != -1 and start_idx <= end_idx:
                    text = text[start_idx:end_idx+1]
                else:
                    app.logger.warning(f"No JSON array brackets found in output. Output was: {text[:100]}...")
                    text = '[]'
                
                # LLM JSON Hatalarını (Trailing Comma vs) temizle
                text = re.sub(r',\s*}', '}', text)
                text = re.sub(r',\s*]', ']', text)
                
                # Truncation Recovery
                if not text.endswith(']'):
                    if text.count('"') % 2 != 0:
                        text += '"'
                    last_brace_idx = text.rfind('}')
                    if last_brace_idx != -1:
                        text = text[:last_brace_idx+1] + '\n]'
                    else:
                        text = '[]'
                
                checklist = json.loads(text)
                break
                
            except Exception as parse_e:
                app.logger.warning(f"Gemini Checklist Hata (Attempt {attempt+1}): {parse_e}")
                time.sleep(2)
                if attempt == max_retries - 1:
                    return jsonify({'error': f'Gemini bağlanılamadı veya çıktı parse edilemedi: {str(parse_e)}'}), 500

        return jsonify(checklist), 200

    except Exception as e:
        app.logger.error(f"Checklist Generate Error: {str(e)}")
        return jsonify({'error': f'Yapay zeka (Checklist) üretimi sırasında hata: {str(e)}'}), 500

@app.route('/api/generate-risks', methods=['POST'])
def generate_risks_ai():
    if not GEMINI_API_KEY:
        return jsonify({'error': 'Sunucuda API Key bulunamadı.'}), 500

    data = request.json
    selected_hazards = data.get('hazards', [])
    context = data.get('context', 'Genel Tesis')
    scope = data.get('scope', 'general') # default: general
    
    if not selected_hazards or not isinstance(selected_hazards, list):
        return jsonify({'error': 'Analiz edilecek tehlike listesi boş olamaz.'}), 400

    system_instruction = """Sen kıdemli bir İş Güvenliği (İSG) Baş Denetçisisin. Sana verilen listedeki HER BİR madde için birbirinden bağımsız, detaylı en az 3-4 farklı risk satırı üreteceksin. Maddeleri asla özetleme veya birleştirme.

KRİTİK GÖREV 1 - DALLANDIRMA: Maddeleri mikro tehlikelere böl. Örn: "Elektrik" -> "Açık kablo", "Kaçak akım rölesi eksikliği", "Islak elle temas".

KRİTİK GÖREV 2 - GERÇEKÇİ PUANLAMA (FINE-KINNEY): Puanlamaları rastgele yapma! Aşağıdaki tabloları baz alarak mesleki tecrübenle mantıklı puanlar seç:

OLASILIK (O):
- 10: Beklenir, kesin (Hemen hemen her gün)
- 6: Yüksek olasılık (Haftada bir)
- 3: Olası (Ayda bir)
- 1: Mümkün fakat düşük (Yılda bir)
- 0.5: Beklenmez (On yılda bir)
- 0.2: Pratik olarak imkansız

FREKANS (F):
- 10: Sürekli (Günde birkaç defa)
- 6: Sık (Günde bir)
- 3: Ara sıra (Haftada bir)
- 2: Sık değil (Ayda bir)
- 1: Seyrek (Yılda birkaç defa)
- 0.5: Çok seyrek (Yılda bir)

ŞİDDET (Ş):
- 100: Çok sayıda ölüm (Katastrofik)
- 40: Ölümlü kaza
- 15: Kalıcı sakatlık, uzuv kaybı (Amputasyon, felç vb.)
- 7: Ciddi yaralanma (Kırık, hastanede yatış)
- 3: Hafif yaralanma (İlkyardım seviyesi, ayakta tedavi)
- 1: Yaralanma yok (Ucuz atlatma)

PUANLAMA KURALLARI:
- Risk Puanı (R) = O x Ş x F.
- Eğer 'riskEffect' kısmında "Ölüm" veya "Vefat" geçiyorsa Şiddet mutlaka 40 veya 100 olmalıdır!
- Eğer "Amputasyon" veya "Kalıcı hasar" geçiyorsa Şiddet mutlaka 15 olmalıdır!
- finalSeverity DAİMA initialSeverity'ye eşit olmalıdır.
- finalProbability ve finalFrequency ilk değerlerin altında olmalıdır.
- KESİNLİKLE YASAK: Hiçbir sayısal değer (initialProbability, initialSeverity, initialFrequency, finalProbability, finalSeverity, finalFrequency) 0 (sıfır) olamaz. En düşük geçerli değer 0.2'dir.

KRİTİK GÖREV 3 - AKTİVİTE ALANI (activity): Her satır için "activity" alanını, o satırdaki spesifik tehlike ve riske göre ayrı ayrı ve profesyonelce doldur. Bu alan "Yapılan İş / Faaliyet / Ekipman / Ortam" sütununa karşılık gelir. ASLA tüm satırlara aynı metni yazma!
Örnekler:
- Tehlike "Açık elektrik kablosu" ise activity: "Elektrik panosu bakım/onarım faaliyeti"
- Tehlike "Forklift çarpması" ise activity: "Forklift ile malzeme taşıma operasyonu"
- Tehlike "Kimyasal temas" ise activity: "Solvent bazlı temizlik işlemi / Kimyasal depolama alanı"
- Tehlike "Gürültü" ise activity: "Kompresör/motor çalışma ortamı – sürekli gürültü maruziyeti"
- Tehlike "Yüksekten düşme" ise activity: "İskele üzerinde çatı onarım faaliyeti"
Her activity metni, o satırdaki tehlike ve bağlamı yansıtmalı; kısa, net ve mesleki bir dille yazılmalıdır.

KRİTİK GÖREV 4 - MEVCUT ÖNLEMLER (existingMeasures): Bu alan için HER ZAMAN ve değişmeden yalnızca şu metni yaz: "Takibi sağlanmaktadır." Başka hiçbir şey yazma.

KRİTİK GÖREV 5 - RİSK / OLASI ETKİ (riskEffect): Bu alanı tıbbi terminoloji kullanarak, klinik düzeyde detaylı doldur. Sadece "yaralanma" veya "kaza" gibi genel ifadeler kullanma. Aşağıdaki örneklere benzer klinik ve patolojik detaylar ekle:
- Mekanik yaralanmalar için: "Kompartman sendromu, açık/kapalı fraktür, tendon/ligaman rüptürü, periferik sinir hasarı"
- Elektrik için: "Ventriküler fibrilasyon, miyokardiyal hasar, derin doku yanığı (dermal/subdermal), rabdomiyoliz"
- Kimyasal/toz için: "Pnömokonyoz (silikoz/antrakoz), obstrüktif akciğer hastalığı (KOAH), kimyasal pnömonit, hipoksemi"
- Gürültü için: "Sensörinöral işitme kaybı (NIHL), tinnitus, koklear hücre hasarı"
- Ergonomi için: "Lomber disk herniasyonu, rotator cuff tendinopatisi, karpal tünel sendromu (median sinir sıkışması)"
- Isı için: "Hipertermik koma, rabdomiyoliz, akut böbrek yetmezliği (ısı çarpmasına bağlı)"
- Yüksekten düşme için: "Politravma, intrakraniyal kanama, vertebral kompresyon fraktürü, dalak/karaciğer laserasyonu"
Her riskEffect cümlesi kısa ama klinik açıdan doğru ve spesifik olmalı; birden fazla patolojiyi virgülle sıralayabilirsin.

KRİTİK GÖREV 6 - DÜZELTİCİ VE ÖNLEYİCİ TEDBİRLER (correctiveMeasures): Bu alan son derece profesyonel, teknik ve standart referanslı olmalıdır. ZORUNLU KURALLAR:
a) Her cümle MUTLAKA "-malıdır" veya "-melidir" eki ile bitmelidir. "-ılmalı", "-iniz", "-ınız" ile bitmez.
b) Teknik standart numaraları, yönetmelik referansları, ölçüm değerleri ve sınır değerleri mutlaka belirtilmelidir.
c) Sayısal eşik değerler kullan: dB(A) değerleri, mg/m³, ppm, lux, °C, kPa, mm, N gibi birimler.
d) Standart/Yönetmelik referansları ekle: TS EN ISO 12100, OSHA 29 CFR 1910, TS EN 60204-1, İSGYÖN, Gürültü Yönetmeliği, Kimyasal Maddelerle Çalışmalarda Sağlık ve Güvenlik Önlemleri Hakkında Yönetmelik, vb.
e) Mühendislik kontrolleri → İdari kontroller → KKD hiyerarşisini takip et.
f) Muayene/test periyotları, kalibrasyon aralıkları, bakım frekanslıları belirtilmelidir.
Örnek kalitede tedbirler:
- "Ortam gürültü düzeyi TS EN ISO 9612 standardına göre en az yılda bir kez ölçülmeli; sekiz saatlik eşdeğer gürültü düzeyi (LEX,8h) 85 dB(A) sınır değerini aşmamalıdır."
- "Tüm döner aksamlara TS EN ISO 14120 standardı kapsamında sabit bariyer koruyucu monte edilmeli; koruyucular araç gerektirmeden sökülemeyecek şekilde emniyete alınmalıdır."
- "Solunabilir kuvars tozu için ACGIH TLV değeri 0,025 mg/m³ aşılmamalı; ortam ölçümü yılda en az bir kez akredite laboratuvara yaptırılmalıdır."
- "Tüm elektrikli ekipmanlar TS EN 60204-1 kapsamında topraklanmalı, kaçak akım röleleri (RCD – 30 mA hassasiyetinde) devre kesicilere entegre edilmelidir."

JSON FORMATI:
Sadece geçerli bir JSON array döndür. Key'ler: activity, hazard, riskEffect, existingMeasures, correctiveMeasures, initialProbability, initialSeverity, initialFrequency, finalProbability, finalSeverity, finalFrequency.
"""

    if scope in ['specific', 'machine']:
        system_instruction += """
DİKKAT KATI KURAL: Kullanıcı senden sadece çok spesifik bir konuyu (veya makineyi) analiz etmeni istiyor. KESİNLİKLE KAPSAM DIŞINA ÇIKMA. Tesisin diğer genel risklerini (eğer kullanıcı özellikle yazmadıysa) ASLA listeye ekleme. Sadece ve sadece kullanıcının girdiği metne %100 sadık kalarak, o dar kapsamdaki mikro tehlikeleri derinlemesine analiz et.
"""
    try:
        all_risks = []
        chunk_size = 5   # Küçük chunk = ~15-20 satır çıktı = hızlı yanıt, timeout yok

        chunks = [selected_hazards[i:i + chunk_size] for i in range(0, len(selected_hazards), chunk_size)]
        
        def process_chunk(chunk, index):
            hazards_text = "\n".join([f"- {h}" for h in chunk])
            prompt_desc = f"Tesis/Bağlam: {context}\nSeçilen Tehlikeler:\n{hazards_text}"

            # Doğrudan Gemini kullan (OpenRouter 402 hatası verdiği için atlandı)
            if not GEMINI_API_KEY:
                return []

            model = genai.GenerativeModel('gemini-2.5-flash', system_instruction=system_instruction)

            max_retries = 3
            for attempt in range(max_retries):
                try:
                    response = model.generate_content(
                        prompt_desc,
                        generation_config=genai.types.GenerationConfig(
                            temperature=0.7,
                            max_output_tokens=8192,
                            response_mime_type="application/json"
                        )
                    )
                    text = response.text.strip()

                    start_idx = text.find('[')
                    end_idx = text.rfind(']')
                    if start_idx != -1 and end_idx != -1 and start_idx <= end_idx:
                        text = text[start_idx:end_idx+1]
                    else:
                        app.logger.warning(f"Chunk {index+1}: JSON array bulunamadı. Çıktı: {text[:200]}")
                        text = '[]'

                    text = re.sub(r',\s*}', '}', text)
                    text = re.sub(r',\s*]', ']', text)

                    # Kesilmiş JSON kurtarma
                    if not text.endswith(']'):
                        if text.count('"') % 2 != 0:
                            text += '"'
                        last_brace_idx = text.rfind('}')
                        if last_brace_idx != -1:
                            text = text[:last_brace_idx+1] + '\n]'
                        else:
                            text = '[]'

                    try:
                        chunk_risks = json.loads(text)
                        if isinstance(chunk_risks, list):
                            app.logger.info(f"Chunk {index+1}/{len(chunks)}: {len(chunk_risks)} risk üretildi.")
                            return chunk_risks
                    except Exception as json_err:
                        app.logger.error(f"Chunk {index+1} JSON parse hatası: {json_err} | Metin: {text[:300]}")

                    return []

                except Exception as e:
                    app.logger.warning(f"Chunk {index+1} Gemini hatası (deneme {attempt+1}/{max_retries}): {e}")
                    if attempt < max_retries - 1:
                        time.sleep(3 * (attempt + 1))  # 3s, 6s backoff
                    else:
                        app.logger.error(f"Chunk {index+1} kalıcı hata: {e}")
            return []

        for idx, chunk in enumerate(chunks):
            try:
                result = process_chunk(chunk, idx)
                if result:
                    all_risks.extend(result)
                    app.logger.info(f"Chunk {idx+1}/{len(chunks)}: {len(result)} risk üretildi.")
            except Exception as e:
                app.logger.error(f"Chunk {idx+1} hatası: {e}")
            if idx < len(chunks) - 1:
                time.sleep(0.5)  # Rate limit baskısını azalt

        return jsonify(all_risks), 200

    except Exception as e:
        app.logger.error(f"Generate Error: {e}")
        return jsonify({'error': f'Yapay zeka matris üretimi sırasında hata: {str(e)}'}), 500

@app.route('/api/dashboard/stats', methods=['GET'])
def get_dashboard_stats():
    """
    Returns global statistics for the premium dashboard.
    """
    total_companies = Company.query.count()
    total_personnel = Personnel.query.count()
    
    # 1. Company Danger Class Distribution
    danger_dist = db.session.query(
        DangerClass.name, func.count(Company.id)
    ).join(NaceCode).join(Company).group_by(DangerClass.name).all()
    
    # 2. Personnel Training Status (Global)
    now = datetime.now().date()
    trained = Personnel.query.filter(Personnel.training_validity_date > now).count()
    expired = Personnel.query.filter(Personnel.training_validity_date <= now).count()
    never_trained = Personnel.query.filter(Personnel.training_validity_date == None).count()
    
    # 3. Health Exam Status (Global)
    healthy = Personnel.query.filter(Personnel.health_validity_date > now).count()
    overdue_health = Personnel.query.filter(Personnel.health_validity_date <= now).count()
    never_exam = Personnel.query.filter(Personnel.health_validity_date == None).count()
    
    # 4. Risk Level Distribution (Global)
    risk_dist = db.session.query(
        RiskAssessment.initial_risk_level, func.count(RiskAssessment.id)
    ).group_by(RiskAssessment.initial_risk_level).all()
    
    # 5. Recent Companies
    recent_companies = Company.query.order_by(Company.created_at.desc()).limit(5).all()
    
    return jsonify({
        'overview': {
            'totalCompanies': total_companies,
            'totalPersonnel': total_personnel,
            'activeProfessionals': CompanyProfessional.query.filter_by(is_active=True).count()
        },
        'dangerDistribution': [{'name': name, 'value': count} for name, count in danger_dist],
        'trainingStats': [
            {'name': 'Eğitimi Tam', 'value': trained},
            {'name': 'Süresi Dolan', 'value': expired},
            {'name': 'Eğitim Almamış', 'value': never_trained}
        ],
        'healthStats': [
            {'name': 'Muayenesi Tam', 'value': healthy},
            {'name': 'Geciken', 'value': overdue_health},
            {'name': 'Henüz Yapılmamış', 'value': never_exam}
        ],
        'riskDistribution': [{'name': name, 'value': count} for name, count in risk_dist],
        'recentCompanies': [{
            'id': c.id,
            'unvan': c.official_title,
            'date': c.created_at.isoformat() if c.created_at else None
        } for c in recent_companies]
    }), 200

# ==========================================
# TRAINING DOCUMENT ROUTES
# ==========================================
ALLOWED_TRAINING_DOC_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg', 'doc', 'docx'}

def allowed_training_doc(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_TRAINING_DOC_EXTENSIONS

@app.route('/api/companies/<int:company_id>/training-documents', methods=['GET'])
def get_training_documents(company_id):
    company = Company.query.get_or_404(company_id)
    docs = TrainingDocument.query.filter_by(company_id=company_id).order_by(TrainingDocument.uploaded_at.desc()).all()
    return jsonify([{
        'id': d.id,
        'docType': d.doc_type.value,
        'docTypeKey': d.doc_type.name,
        'originalFilename': d.original_filename,
        'uploadedAt': d.uploaded_at.isoformat() if d.uploaded_at else None,
        'year': (d.uploaded_at or datetime.utcnow()).year,
    } for d in docs]), 200

def _resolve_training_doc_path(upload_dir, stored_filename):
    """Returns the full path to a training document file, with legacy fallback."""
    filepath = os.path.join(upload_dir, stored_filename)
    if not os.path.exists(filepath):
        # Legacy: file may be in flat directory (no year subfolder)
        flat = os.path.join(upload_dir, os.path.basename(stored_filename))
        if os.path.exists(flat):
            return flat
    return filepath

def _migrate_training_docs_to_year_folders():
    """Move existing flat-stored training docs into year subfolders."""
    upload_dir = app.config['TRAINING_DOCS_FOLDER']
    docs = TrainingDocument.query.filter(~TrainingDocument.stored_filename.contains('/')).all()
    for doc in docs:
        flat_path = os.path.join(upload_dir, doc.stored_filename)
        if not os.path.exists(flat_path):
            continue
        year = (doc.uploaded_at or datetime.utcnow()).year
        year_dir = os.path.join(upload_dir, str(year))
        os.makedirs(year_dir, exist_ok=True)
        new_stored = f"{year}/{doc.stored_filename}"
        new_path = os.path.join(upload_dir, new_stored)
        os.rename(flat_path, new_path)
        doc.stored_filename = new_stored
    db.session.commit()

@app.route('/api/companies/<int:company_id>/training-documents', methods=['POST'])
def upload_training_document(company_id):
    company = Company.query.get_or_404(company_id)

    if 'file' not in request.files:
        return jsonify({'error': 'Dosya bulunamadı'}), 400

    file = request.files['file']
    doc_type_key = request.form.get('docType')

    if not file or not file.filename:
        return jsonify({'error': 'Dosya seçilmedi'}), 400

    if not allowed_training_doc(file.filename):
        return jsonify({'error': 'Desteklenmeyen dosya formatı. İzin verilen: PDF, PNG, JPG, DOC, DOCX'}), 400

    # Validate doc type
    try:
        dtype = TrainingDocTypeEnum[doc_type_key]
    except (KeyError, TypeError):
        return jsonify({'error': 'Geçersiz belge türü'}), 400

    upload_dir = app.config['TRAINING_DOCS_FOLDER']
    year = datetime.utcnow().year
    year_dir = os.path.join(upload_dir, str(year))
    os.makedirs(year_dir, exist_ok=True)

    original_name = file.filename
    safe_name = secure_filename(file.filename)
    base_stored = f"{company_id}_{int(datetime.utcnow().timestamp())}_{safe_name}"
    stored_name = f"{year}/{base_stored}"
    file.save(os.path.join(upload_dir, stored_name))

    doc = TrainingDocument(
        company_id=company_id,
        doc_type=dtype,
        original_filename=original_name,
        stored_filename=stored_name,
    )
    db.session.add(doc)
    db.session.commit()

    return jsonify({
        'id': doc.id,
        'docType': doc.doc_type.value,
        'docTypeKey': doc.doc_type.name,
        'originalFilename': doc.original_filename,
        'uploadedAt': doc.uploaded_at.isoformat() if doc.uploaded_at else None,
        'message': 'Belge yüklendi'
    }), 201

@app.route('/api/training-documents/<int:doc_id>/download')
def download_training_document(doc_id):
    doc = TrainingDocument.query.get_or_404(doc_id)
    upload_dir = app.config['TRAINING_DOCS_FOLDER']
    filepath = _resolve_training_doc_path(upload_dir, doc.stored_filename)
    if not os.path.exists(filepath):
        return jsonify({'error': 'Dosya bulunamadı'}), 404
    return send_file(filepath, download_name=doc.original_filename, as_attachment=True)

@app.route('/api/training-documents/<int:doc_id>/view')
def view_training_document(doc_id):
    doc = TrainingDocument.query.get_or_404(doc_id)
    upload_dir = app.config['TRAINING_DOCS_FOLDER']
    filepath = _resolve_training_doc_path(upload_dir, doc.stored_filename)
    if not os.path.exists(filepath):
        return jsonify({'error': 'Dosya bulunamadı'}), 404
    return send_file(filepath, download_name=doc.original_filename, as_attachment=False)

@app.route('/api/training-documents/<int:doc_id>', methods=['DELETE'])
def delete_training_document(doc_id):
    doc = TrainingDocument.query.get_or_404(doc_id)
    upload_dir = app.config['TRAINING_DOCS_FOLDER']
    filepath = _resolve_training_doc_path(upload_dir, doc.stored_filename)
    if os.path.exists(filepath):
        os.remove(filepath)
    db.session.delete(doc)
    db.session.commit()
    return jsonify({'message': 'Belge silindi'}), 200

# ==========================================
# SAĞLIK BELGELERİ (Health Documents)
# ==========================================
ALLOWED_HEALTH_DOC_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg', 'doc', 'docx'}

def allowed_health_doc(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_HEALTH_DOC_EXTENSIONS

@app.route('/api/companies/<int:company_id>/health-documents', methods=['GET'])
def get_health_documents(company_id):
    Company.query.get_or_404(company_id)
    docs = HealthDocument.query.filter_by(company_id=company_id).order_by(HealthDocument.uploaded_at.desc()).all()
    return jsonify([{
        'id': d.id,
        'docType': d.doc_type.value,
        'docTypeKey': d.doc_type.name,
        'originalFilename': d.original_filename,
        'uploadedAt': d.uploaded_at.isoformat() if d.uploaded_at else None,
        'year': (d.uploaded_at or datetime.utcnow()).year,
    } for d in docs]), 200

@app.route('/api/companies/<int:company_id>/health-documents', methods=['POST'])
def upload_health_document(company_id):
    Company.query.get_or_404(company_id)

    if 'file' not in request.files:
        return jsonify({'error': 'Dosya bulunamadı'}), 400

    file = request.files['file']
    doc_type_key = request.form.get('docType')

    if not file or not file.filename:
        return jsonify({'error': 'Dosya seçilmedi'}), 400

    if not allowed_health_doc(file.filename):
        return jsonify({'error': 'Desteklenmeyen dosya formatı. İzin verilen: PDF, PNG, JPG, DOC, DOCX'}), 400

    try:
        dtype = HealthDocTypeEnum[doc_type_key]
    except (KeyError, TypeError):
        return jsonify({'error': 'Geçersiz belge türü'}), 400

    upload_dir = app.config['HEALTH_DOCS_FOLDER']
    year = datetime.utcnow().year
    year_dir = os.path.join(upload_dir, str(year))
    os.makedirs(year_dir, exist_ok=True)

    original_name = file.filename
    safe_name = secure_filename(file.filename)
    base_stored = f"{company_id}_{int(datetime.utcnow().timestamp())}_{safe_name}"
    stored_name = f"{year}/{base_stored}"
    file.save(os.path.join(upload_dir, stored_name))

    doc = HealthDocument(
        company_id=company_id,
        doc_type=dtype,
        original_filename=original_name,
        stored_filename=stored_name,
    )
    db.session.add(doc)
    db.session.commit()

    return jsonify({
        'id': doc.id,
        'docType': doc.doc_type.value,
        'docTypeKey': doc.doc_type.name,
        'originalFilename': doc.original_filename,
        'uploadedAt': doc.uploaded_at.isoformat() if doc.uploaded_at else None,
        'message': 'Belge yüklendi'
    }), 201

@app.route('/api/health-documents/<int:doc_id>/download')
def download_health_document(doc_id):
    doc = HealthDocument.query.get_or_404(doc_id)
    upload_dir = app.config['HEALTH_DOCS_FOLDER']
    filepath = _resolve_training_doc_path(upload_dir, doc.stored_filename)
    if not os.path.exists(filepath):
        return jsonify({'error': 'Dosya bulunamadı'}), 404
    return send_file(filepath, download_name=doc.original_filename, as_attachment=True)

@app.route('/api/health-documents/<int:doc_id>/view')
def view_health_document(doc_id):
    doc = HealthDocument.query.get_or_404(doc_id)
    upload_dir = app.config['HEALTH_DOCS_FOLDER']
    filepath = _resolve_training_doc_path(upload_dir, doc.stored_filename)
    if not os.path.exists(filepath):
        return jsonify({'error': 'Dosya bulunamadı'}), 404
    return send_file(filepath, download_name=doc.original_filename, as_attachment=False)

@app.route('/api/health-documents/<int:doc_id>', methods=['DELETE'])
def delete_health_document(doc_id):
    doc = HealthDocument.query.get_or_404(doc_id)
    upload_dir = app.config['HEALTH_DOCS_FOLDER']
    filepath = _resolve_training_doc_path(upload_dir, doc.stored_filename)
    if os.path.exists(filepath):
        os.remove(filepath)
    db.session.delete(doc)
    db.session.commit()
    return jsonify({'message': 'Belge silindi'}), 200

# ==========================================
# BİLDİRİMLER (Notifications)
# ==========================================
@app.route('/api/notifications', methods=['GET'])
def get_notifications():
    from datetime import date, timedelta
    today = date.today()
    thirty_days = today + timedelta(days=30)
    notifications = []

    companies = Company.query.filter_by(is_active=True).all()

    for c in companies:
        name = c.short_name or c.official_title

        # 1. İSG Profesyoneli atanmamış firmalar
        active_profs = [p for p in c.professionals if p.is_active]
        has_igu = any(p.role == RoleEnum.IGU for p in active_profs)
        has_iyh = any(p.role == RoleEnum.IYH for p in active_profs)
        if not has_igu:
            notifications.append({
                'icon': 'warning',
                'type': 'danger',
                'text': f'{name} firmasına İş Güvenliği Uzmanı (İGU) atanmamış.',
                'time': 'Acil'
            })
        if not has_iyh:
            notifications.append({
                'icon': 'warning',
                'type': 'danger',
                'text': f'{name} firmasına İşyeri Hekimi (İYH) atanmamış.',
                'time': 'Acil'
            })

        # 2. Çalışan sayısı 0 olan firmalar
        if c.employee_count == 0:
            notifications.append({
                'icon': 'group_off',
                'type': 'warning',
                'text': f'{name} firmasında kayıtlı çalışan bulunmuyor.',
                'time': 'Bilgi'
            })

        # 3. Eğitim ve sağlık gözetimi süresi dolmuş / dolmak üzere personeller
        for p in c.personnel:
            full = f'{p.first_name} {p.last_name}'
            if p.training_validity_date:
                if p.training_validity_date < today:
                    notifications.append({
                        'icon': 'school',
                        'type': 'danger',
                        'text': f'{name} — {full} eğitim geçerlilik süresi dolmuş.',
                        'time': 'Süresi geçti'
                    })
                elif p.training_validity_date <= thirty_days:
                    notifications.append({
                        'icon': 'school',
                        'type': 'warning',
                        'text': f'{name} — {full} eğitim süresi 30 gün içinde dolacak.',
                        'time': 'Yaklaşıyor'
                    })
            elif p.is_active:
                notifications.append({
                    'icon': 'school',
                    'type': 'info',
                    'text': f'{name} — {full} için eğitim kaydı bulunmuyor.',
                    'time': 'Eksik'
                })

            if p.health_validity_date:
                if p.health_validity_date < today:
                    notifications.append({
                        'icon': 'health_and_safety',
                        'type': 'danger',
                        'text': f'{name} — {full} sağlık muayenesi süresi dolmuş.',
                        'time': 'Süresi geçti'
                    })
                elif p.health_validity_date <= thirty_days:
                    notifications.append({
                        'icon': 'health_and_safety',
                        'type': 'warning',
                        'text': f'{name} — {full} sağlık muayenesi 30 gün içinde dolacak.',
                        'time': 'Yaklaşıyor'
                    })

        # 4. Risk değerlendirmesi belgesi kontrol
        risk_docs = [d for d in c.documents if d.document_type == DocumentTypeEnum.Risk_Degerlendirmesi]
        if not risk_docs:
            notifications.append({
                'icon': 'assignment_late',
                'type': 'danger',
                'text': f'{name} firmasının risk değerlendirmesi belgesi yok.',
                'time': 'Eksik'
            })
        else:
            latest_risk = max(risk_docs, key=lambda d: d.revision_date)
            if latest_risk.revision_date < today:
                notifications.append({
                    'icon': 'assignment_late',
                    'type': 'warning',
                    'text': f'{name} risk değerlendirmesi revizyon tarihi geçmiş.',
                    'time': 'Güncellenmeli'
                })

        # 5. Acil durum planı eksik
        adp_docs = [d for d in c.documents if d.document_type == DocumentTypeEnum.Acil_Durum_Plani]
        if not adp_docs:
            notifications.append({
                'icon': 'local_fire_department',
                'type': 'warning',
                'text': f'{name} firmasının acil durum planı eksik.',
                'time': 'Eksik'
            })

    # Sıralama: danger > warning > info
    priority = {'danger': 0, 'warning': 1, 'info': 2}
    notifications.sort(key=lambda n: priority.get(n['type'], 3))

    return jsonify(notifications[:50]), 200


setup_db()

if __name__ == '__main__':
    with app.app_context():
        _migrate_training_docs_to_year_folders()
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('FLASK_ENV', 'development') != 'production'
    app.run(debug=debug, host='0.0.0.0', port=port)
